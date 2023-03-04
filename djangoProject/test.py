import base64
import datetime
import json
import os
import re
import time
import urllib.parse

import requests
import pandas as pd
from docx import Document
from openpyxl.styles import Border, Side, Alignment, Font

from sqlalchemy import create_engine

from djangoProject.color import mos_wirte_excel

engine = create_engine("mysql+pymysql://root:123456@127.0.0.1:3306/wubian")

from djangoProject.erpApi import requesterp
from djangoProject.lingxingApi.aes import md5_encrypt, aes_encrypt
import orjson
from djangoProject.lingxingApi.http_util import HttpBase
from djangoProject.lingxingApi.resp_schema import ResponseResult
from djangoProject.lingxingApi.sign import SignBase
# a = requesterp()
def listing_price_load():
    df = pd.read_excel(r"C:\Users\wb\Desktop\Listing20230303-490113876584333312.xlsx").fillna("")
    df = df.rename(columns={"店铺":"store","ASIN":"asin","父ASIN":"parent_asin","MSKU":"msku","SKU":"GS-sku","价格":"price"},inplace=False)
    df = df[['store','asin','parent_asin','msku','GS-sku','price']]
    df["date"] = datetime.datetime.now()
    df["sku"] = df["GS-sku"].replace("GS-","")
    df["update_date"] = datetime.datetime.now()
    df["create_date"] = datetime.datetime.now()
    df.to_sql('listing_price_date',engine,chunksize=10000,if_exists='append',index=False)
    print("done")


# listing_price_load()


def api_listing_price_load():
    df_sid = pd.DataFrame()
    sids = requests.get("http://43.142.117.35/get_store_id")
    data = eval(sids.text)
    for i in data:
        sid = i.get("sid")
        name = i.get("name")
        dftmp = pd.DataFrame({"sid":sid,"name":name},index=[0])
        df_sid = df_sid.append(dftmp)
    df_sid = df_sid.reset_index()
    del df_sid["index"]
    offset = 0
    df_all = pd.DataFrame()
    while True:
        res = requests.get("http://43.142.117.35/get_listing",params={"offset":offset}).json()
        data = res.get("data")
        df = pd.DataFrame()

        for i in data:
            status = i.get("status")
            is_delete = i.get("is_delete")
            if is_delete == 0 and status == 1:
                asin = i.get("asin")
                parent_asin = i.get("parent_asin")
                msku = i.get("seller_sku")
                sid = i.get("sid")
                GS_sku = i.get("local_sku")
                sku = GS_sku.replace("GS-","")
                price = i.get("price")
                dftmp = pd.DataFrame({"sid":sid,"asin":asin,"parent_asin":parent_asin,"msku":msku,"status":status,"GS-sku":GS_sku,"sku":sku,"price":price,"is_delete":is_delete},index=[0])
                df = df.append(dftmp)
        df = df.reset_index()
        del df['index']
        df = pd.merge(df,df_sid,left_on=['sid'],right_on=['sid'],how="left").reset_index()
        del df['index']
        df["date"] = datetime.datetime.now()
        df["update_date"] = datetime.datetime.now()
        df["create_date"] = datetime.datetime.now()
        df = df.rename(
            columns={"name": "store"},
            inplace=False)
        df = df[['store', 'asin', 'parent_asin', 'msku', 'GS-sku','sku','price','date','update_date','create_date']]
        df_all.append(df)
    df_all.to_sql('listing_price_date', engine, chunksize=10000, if_exists='append', index=False)
        # df.to_excel(r"C:\Users\wb\Desktop\listing.xlsx")


# api_listing_price_load()
# start_date = '2023-01-01'
# end_date = '2023-01-31'
# req = {
#     "start_date":start_date,
#     "end_date":end_date,
# }
# a = requests.post("http://43.142.117.35/old_profitSettlement",params=req)
# a=11
req_body = {
    "sid": 2613,
    "offset": 0,
    "length": 1000,
    "sort_type": "desc",
    "sort_field": "volume",
    "start_date": "2023-02-13",
    "end_date": "2023-02-13",
    "summary_field": "msku",
    }

# a = requests.post("http://43.142.117.35/get_fbawarehouse")
# a = requests.post("http://43.142.117.35/productPerFormance",params=req_body)
# b = eval(json.dumps(a.text))
# eval(b)
# a.json()
a=1

#
def 箱标():
    import winreg
    def desktop_path():
        key = winreg.OpenKey(winreg.HKEY_CURRENT_USER,
                             r'Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders')
        desktop = winreg.QueryValueEx(key, "Desktop")[0]
        print(desktop)
        return desktop
    from docx import Document
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
    desktop_path = desktop_path()
    doc = Document(desktop_path+"\日本海外仓箱标\日本海外仓箱标模板.docx")
    df = pd.read_excel(desktop_path+"\日本海外仓箱标\需生成的数据文件 - 副本.xlsx")
    df_data = pd.DataFrame()
    for i in df.index:
        xiangshu = int(df.loc[i,"箱数"])
        sku = df.loc[i,"品名"]
        num = int(df.loc[i,"单箱数量"])
        zimu = df.loc[i,"字母标签"]
        code = df.loc[i,"箱子编号"]
        if xiangshu>1:
            df_tmp = pd.DataFrame({"品名":sku,"箱子编号":code,'字母标签':zimu,"单箱数量":num},index=[0])
            for n in range(0,xiangshu):
                df_data = df_data.append(df_tmp)
        else:
            df_tmp = pd.DataFrame({"品名": sku, "箱子编号": code, '字母标签': zimu, "单箱数量": num}, index=[0])
            df_data = df_data.append(df_tmp)
    df_data = df_data.reset_index()
    sku_list = df_data['品名'].tolist()
    num_list = df_data['单箱数量'].tolist()
    zimu_list = df_data['字母标签'].tolist()
    code_list = df_data['箱子编号'].tolist()
    list1 = doc.paragraphs
    for i in list1:
        delete_paragraph(i)
        if (len(doc.paragraphs))==len(sku_list):
            break
    # doc.save(desktop_path+"\日本海外仓箱标\日本海外仓箱标.docx")
    doc.save(r'C:\Users\wb\Desktop\test.docx')

    doc = Document(r'C:\Users\wb\Desktop\test.docx')
    children = doc.element.body.iter()
    count = 0  # 写一个count是为了，可以定位是哪个文本框，因为我用索引失败了
    for child in children:
        # 通过类型判断目录
        if child.tag.endswith('txbx'):
            count += 1
            if count == 3:
                for ci in child.iter():
                    if ci.tag.endswith('main}r'):
                        if ci.text == '型号：':
                            ci.text = ''
                        if ci.text == '字母区分：':
                            ci.text = ''
                        if ci.text == '数量：':
                            ci.text = ''
                        if ci.text == '箱子编号：':
                            ci.text = ''
    doc.save(r'C:\Users\wb\Desktop\test.docx')
    doc = Document(r'C:\Users\wb\Desktop\test.docx')
    sku_dict = {}
    for i in range(1,(len(sku_list)*2)+1):
        if i==1:
            sku_dict[1] = 1
        elif i%2==0:
            sku_dict[i] = sku_dict[i-1]
        else:
            sku_dict[i] = sku_dict[i-1] + 1
    m = 1
    n = 0
    children = doc.element.body.iter()
    for child in children:
        # 通过类型判断目录
        if child.tag.endswith('txbx'):
            for ci in child.iter():
                if ci.tag.endswith('main}r'):
                    if ci.text == '型号：':
                        # print(ci.text)
                        ci.text = "型号：%s"%sku_list[n]
                    if ci.text == '字母区分：':
                        # print(ci.text)
                        ci.text = "字母区分：%s" % zimu_list[n]
                    if ci.text == '数量：':
                        # print(ci.text)
                        ci.text = "数量：%s" % num_list[n]
                    if ci.text == '箱子编号：':
                        # print(ci.text)
                        ci.text = "箱子编号：\n%s"%code_list[n]
                        n+=1
                print("正在生成")
        doc.save(r'C:\Users\wb\Desktop\test.docx')

#

# 箱标()

# for index,sku in enumerate(sku_list):




def 补货表生成V1():

    #sales and traffic 数据透视处理
    df_7D = pd.read_excel(r'C:\Users\wb\Desktop\工作簿2.xlsx',sheet_name="7天")
    df_7D = df_7D.pivot_table(index=['（子）ASIN'],values=['会话次数 – 总计','会话次数 – 总计 – B2B','订单商品总数','订单商品总数 - B2B'],aggfunc=sum).reset_index()
    df_14D = pd.read_excel(r'C:\Users\wb\Desktop\工作簿2.xlsx',sheet_name="14天")
    df_14D = df_14D.pivot_table(index=['（子）ASIN'],values=['会话次数 – 总计','会话次数 – 总计 – B2B','订单商品总数','订单商品总数 - B2B'],aggfunc=sum).reset_index()
    df_magInv = pd.read_excel(r'C:\Users\wb\Desktop\工作簿2.xlsx',sheet_name="管理亚马逊库存")
    df_asin_to_sku = df_magInv[['sku','asin']]
    df_7D = pd.merge(df_7D,df_asin_to_sku,left_on=['（子）ASIN'],right_on=['asin'],how='left')
    df_14D = pd.merge(df_14D,df_asin_to_sku,left_on=['（子）ASIN'],right_on=['asin'],how='left')
    del df_7D["asin"]
    del df_14D["asin"]
    df_7D['总流量'] = df_7D['会话次数 – 总计'] + df_7D['会话次数 – 总计 – B2B']
    df_7D['总出单'] = df_7D['订单商品总数'] + df_7D['订单商品总数 - B2B']
    df_7D['转化率'] = df_7D['总出单'] / df_7D['总流量']
    df_14D['总流量'] = df_14D['会话次数 – 总计'] + df_14D['会话次数 – 总计 – B2B']
    df_14D['总出单'] = df_14D['订单商品总数'] + df_14D['订单商品总数 - B2B']
    df_14D['转化率'] = df_14D['总出单'] / df_14D['总流量']
    df_7D['转化率'] = df_7D['转化率'].map(lambda x: format(x,'.2%'))
    df_14D['转化率'] = df_14D['转化率'].map(lambda x: format(x,'.2%'))



    #补货表数据处理
    df_buhuo = pd.read_excel(r'C:\Users\wb\Desktop\工作簿2.xlsx',sheet_name="补货")
    df_buhuo['预留'] = df_buhuo['FC transfer'] + df_buhuo['FC Processing']

    #管理亚马逊库存
    df_magInv = pd.merge(df_magInv,df_buhuo[['ASIN','预留']],left_on=['asin'],right_on=['ASIN'],how="left")
    df_magInv['IN-STOCK库存'] = df_magInv['afn-fulfillable-quantity']
    df_magInv['总库存'] = df_magInv['预留'] + df_magInv['afn-fulfillable-quantity'] +df_magInv['afn-inbound-working-quantity'] + df_magInv['afn-inbound-shipped-quantity'] + df_magInv['afn-inbound-receiving-quantity']


    #库存明细生成
    df_Invdetail = df_asin_to_sku
    df_Invdetail.rename(columns={"sku": "MSKU"},inplace=True)
    df_basedata = pd.read_excel(r'C:\Users\wb\Desktop\补货表基础信息.xlsx',sheet_name="基础信息")
    df_Invdetail = pd.merge(df_Invdetail,df_basedata,
                            left_on=['MSKU'],right_on=['MSKU'],how="outer").fillna(0)


    df_Invdetail = pd.merge(df_Invdetail,df_magInv[['sku','IN-STOCK库存','总库存','your-price']],left_on=['MSKU'],right_on=['sku'],how="left").fillna("")
    df_Invdetail.rename(columns={"总库存": "亚马逊总库存"},inplace=True)
    df_Invdetail['FBA在库预估'] = df_Invdetail['亚马逊总库存']
    df_Invdetail['海外仓库存'] = 0
    df_Invdetail['预创货件'] = 0
    df_Invdetail = pd.merge(df_Invdetail,df_7D[['sku','总流量','转化率','总出单']],left_on=['sku'],right_on=['sku'],how="left")
    df_Invdetail.rename(columns={"总流量": "7天流量","转化率":"7天转化率","总出单":"7天订单数量"},inplace=True)
    df_Invdetail['7天实际日均'] = round(df_Invdetail['7天订单数量']/7,1)
    df_Invdetail['常规补货计划日均'] = 0
    df_Invdetail = pd.merge(df_Invdetail,df_14D[['sku','总流量','转化率','总出单']],left_on=['sku'],right_on=['sku'],how="left")
    df_Invdetail.rename(columns={"总流量": "14天流量","转化率":"14天转化率","总出单":"14天订单数量"},inplace=True)
    df_Invdetail['14天日均数量'] = round(df_Invdetail['14天订单数量']/14)
    df_Invdetail['7/14增长'] = (df_Invdetail['7天实际日均']-df_Invdetail['14天日均数量'])/df_Invdetail['14天日均数量']
    df_Invdetail.rename(columns={"asin": "ASIN","your-price":"价格"},inplace=True)
    df_Invdetail = df_Invdetail.fillna(0)
    weight_av = 6000
    rate = 6.6
    canshu = 7
    df_Invdetail['计费重'] = round((df_Invdetail['长']*df_Invdetail['宽']*df_Invdetail['高'])/weight_av)
    df_Invdetail['价格'] = df_Invdetail['价格'].replace("",0)
    df_Invdetail["$头程"] = round(df_Invdetail['计费重']*canshu/rate,1)

    df_Invdetail['分割1'] = ""
    df_Invdetail['分割2'] = ""
    df_Invdetail['3'] = ""
    df_Invdetail['4'] = ""
    df_Invdetail['5'] = ""
    df_Invdetail['6'] = ""


    df_Invdetail = df_Invdetail[["MSKU","公司SKU","数据","父体","X备注","备注","分割1","颜色大类",
                                 "细分颜色","组合方式","材质","可扩展",
                                 "20寸的尺寸","ASIN","分割2","IN-STOCK库存","FBA在库预估",
                                 "亚马逊总库存","3","海外仓库存","预创货件","4","7天流量","7天转化率",
                                 "7天订单数量","7天实际日均","常规补货计划日均","5","14天订单数量",
                                 "14天日均数量","7/14增长","14天流量","14天转化率","6","佣金","FBA费",
                                 "成本","高","长","宽","计费重","$头程","价格"]]
    df_Invdetail['定价毛利率'] = ((df_Invdetail['价格']-df_Invdetail['价格']*df_Invdetail['佣金']-df_Invdetail['FBA费']-(df_Invdetail['成本'])/rate)-df_Invdetail['$头程'])/df_Invdetail['价格']
    df_Invdetail['DEAL'] = 0
    df_Invdetail['Coupon'] = 0
    df_Invdetail['Prime'] = 0
    df_Invdetail['OFF'] = 0
    df_Invdetail['售价'] = df_Invdetail['价格'] * (1-df_Invdetail['OFF'])-df_Invdetail['DEAL']*df_Invdetail['价格']-df_Invdetail['Coupon']-df_Invdetail['Prime']*df_Invdetail['价格']
    df_Invdetail['毛利率'] = (df_Invdetail['售价']-df_Invdetail['售价']*df_Invdetail['佣金']-df_Invdetail['成本']/rate-df_Invdetail['FBA费']-df_Invdetail['$头程'])/df_Invdetail['售价']
    df_Invdetail['毛利额$'] = round(df_Invdetail['售价'] * df_Invdetail['毛利率'],2)
    df_Invdetail['7'] = ""
    df_Invdetail['INSTOCK库存售罄天数'] = "=IFERROR(P2/AA2,0)"
    df_Invdetail['FBA预估库存售罄天数'] = "=IFERROR(Q2/AA2,0)"
    df_Invdetail['FBA总库存（含在途）售罄天数'] = "=IFERROR(R2/AA2,0)"
    df_Invdetail['INSTOCK差额＜60'] = "=BC2-BB2"
    df_Invdetail['实际售罄天数'] = "=IFERROR(R2/Z2,0)"
    df_Invdetail['8'] = ""
    df_Invdetail['9'] = ""
    df_Invdetail['安全天数'] = ""
    df_Invdetail['应补货'] = "=IF(AA2*BH2-R2>0,AA2*BH2-R2,0)"
    df_Invdetail['备注2'] = ""
    df_Invdetail['1月待安排'] = ""
    df_Invdetail['10'] = ""
    df_Invdetail["15天特批"] = ""
    df_Invdetail["25天"] = ""
    df_Invdetail["35天"] = ""
    df_Invdetail["45天"] = ""
    df_Invdetail["60天"] = ""
    df_Invdetail["合计"] = "=SUM(BM2:BQ2)"
    df_Invdetail["复核天数"] = "=IFERROR(BR2/AA2,0)"
    df_Invdetail["合计天数"] = "=IFERROR(BC2+BS2,0)"
    df_Invdetail["差值"] = "=IFERROR(BR2-BI2,0)"
    df_Invdetail['11'] = ""
    df_Invdetail["工厂待出货数量"] = ""
    df_Invdetail["预下单量"] = ""
    df_Invdetail['12'] = ""
    df_Invdetail["x备注"] = ""
    df_Invdetail["F"] = ""
    df_Invdetail['定价毛利率'] = df_Invdetail['定价毛利率'].map(lambda x: format(x,'.2%'))
    df_Invdetail['佣金'] = df_Invdetail['佣金'].map(lambda x: format(x,'.2%'))
    writer = pd.ExcelWriter(r"C:\Users\wb\Desktop\LXGG.xlsx", engine='openpyxl')  # 创建数据存放路径
    df_Invdetail.to_excel(writer, index=False, sheet_name='Sheet1')
    writer.save()  # 文件保存
    writer.close()

    a=11



# 补货表生成V1()

# a=11


def 供应商问题标签():
    df = pd.read_excel(r'C:\Users\wb\Desktop\拉杆箱次品原因汇总.xlsx',sheet_name="Sheet2")

    for i in df.index:
        list1 = df.loc[i,'汇总'].split("，")
        list_name = df.columns.values.tolist()
        for j in list1:
            if j.strip() not in list_name:
                df.loc[i,'%s'%j.strip()] = 1
            else:
                df.loc[i, '%s' % j.strip()] = 1
        a=1

    df=df.fillna(0)
    df.to_excel(r"C:\Users\wb\Desktop\123.xlsx")
    a=11



    import re
    df = pd.read_excel(r'C:\Users\wb\Desktop\拉杆箱次品原因汇总.xlsx',sheet_name="Sheet1")
    for i in df.index:
        list1 = df.loc[i,'汇总'].split("，")
        list_name = df.columns.values.tolist()
        for j in list1:
            if j.strip() not in list_name:
                if j == "内装折皱批量":
                    df.loc[i, '内装折皱'] = "批量"
                elif j == "三角不平批量":
                    df.loc[i, '三角不平'] = "批量"
                else:
                    try:
                        num = re.search("\d+",j)
                        num = int(num.group(0))
                        name = j.replace("%s个"%num,"")
                        name = name.replace("个%s" %num,"").strip()
                        df.loc[i,'%s'%name] = num
                    except:
                        df.loc[i, '%s'%j.strip()] = 1
            else:
                try:
                    num = re.search("\d+", j)
                    num = int(num.group(0))
                    name = j.replace("%s个" % num, "")
                    name = name.replace("个%s" % num, "")
                    df.loc[i, '%s' % name] = num
                except:
                    pass

        a=1

    df.to_excel(r"C:\Users\wb\Desktop\234.xlsx")



    df_1 = pd.read_excel(r'C:\Users\wb\Desktop\123.xlsx')
    df_2 = pd.read_excel(r'C:\Users\wb\Desktop\234.xlsx')

    list1 = df_1.columns.values.tolist()
    list2 = df_2.columns.values.tolist()



    df = df_1.append(df_2)
    df.to_excel(r"C:\Users\wb\Desktop\aaaa.xlsx")

    #
    # a=11

#
#
def write_excel():
    Inv_arr = ["MSKU","公司SKU","数据","父体","X备注","备注","分割1","颜色大类",
                                     "细分颜色","组合方式","材质","可扩展",
                                     "20寸的尺寸","ASIN","分割2","IN-STOCK库存","FBA在库预估",
                                     "亚马逊总库存","3","海外仓库存","预创货件","4","7天流量","7天转化率",
                                     "7天订单数量","7天实际日均","常规补货计划日均","5","14天订单数量",
                                     "14天日均数量","7/14增长","14天流量","14天转化率","6","佣金","FBA费",
                                     "成本","高","长","宽","计费重","$头程","价格","定价毛利率","DEAL",
                "Coupon","Prime","OFF","售价","毛利率","毛利额$","7","定价毛利率",
               "INSTOCK库存售罄天数","FBA预估库存售罄天数",
               "FBA总库存（含在途）售罄天数","INSTOCK差额＜60",
               "实际售罄天数","8","9","安全天数","应补货","备注2","1月待安排","10","备注"
               "15天特批","25天","35天","45天","60天","复核天数",
               "合计天数","差值","11","工厂待出货数量","预下单量","12","x备注","F"]
    inv_dic = {}
    title = Inv_arr
    import openpyxl
    excel=openpyxl.load_workbook(r'C:\Users\wb\Desktop\LXGG.xlsx')
    sheet=excel['Sheet1']
    sheet.freeze_panes='AB2' #冻结单元格
    zm = [chr(i) for i in range(ord("A"), ord("Z") + 1)]
    for t in title:
        if title.index(t) == 26:
            zm.extend(["A%s" % i for i in zm])
        if title.index(t) == 53:
            zm.extend(["B%s" % i for i in zm])
        sheet[f"{zm[title.index(t)]}1"].value = t
        inv_dic.update({t: zm[title.index(t)]})  # {"标题": 列序号}


    border = Border(left=Side(border_style='thin', color='000000'),
                    right=Side(border_style='thin', color='000000'),
                    top=Side(border_style='thin', color='000000'),
                    bottom=Side(border_style='thin', color='000000'))

    #列染色和设置分割列宽,type=1为分割列
    def col_colors(names,color,type=0):
        from openpyxl.styles import PatternFill
        for name in names:
            col = sheet[inv_dic['%s'%name]]
            fills = PatternFill("solid", fgColor=color)
            if type == 1:
                sheet.column_dimensions[inv_dic["%s"%name]].width = 0.46
            for cell in col:
                cell.fill = fills

    sheet.column_dimensions.group(inv_dic['公司SKU'],inv_dic['父体'], hidden=True)
    sheet.column_dimensions.group(inv_dic['颜色大类'],inv_dic['ASIN'], hidden=True)
    sheet.column_dimensions.group(inv_dic['佣金'],inv_dic['$头程'], hidden=True)
    sheet.column_dimensions.group(inv_dic['DEAL'],inv_dic['毛利额$'], hidden=True)
    sheet.column_dimensions.group(inv_dic['备注2'],inv_dic['1月待安排'], hidden=True)
    # sheet.column_dimensions.outline_level=None

    col_colors(['分割1','分割2','3','4','5','6','7','8','9','10','11','12'],"C00000",type=1)
    col_colors(['FBA在库预估'],"B4C6E7")
    col_colors(['海外仓库存','预创货件'],"FFF2CC")
    col_colors(['7天流量','7天转化率'],"D6DCE4")
    col_colors(['常规补货计划日均'],"F8CBAD")
    col_colors(['INSTOCK库存售罄天数','FBA预估库存售罄天数','FBA总库存（含在途）售罄天数','INSTOCK差额＜60'],"D6DCE4")
    col_colors(["14天流量","14天转化率","佣金","FBA费","成本","高","长","宽","计费重","$头程"],"D6DCE4")
    col_colors(["价格","定价毛利率"],"E2EFDA")
    col_colors(["实际售罄天数"],"FCE4D6")
    col_colors(["安全天数"],"DDEBF7")
    col_colors(["应补货"],"F4B084")
    col_colors(["备注2","1月待安排"],"FFF2CC")
    col_colors(["25天","35天"],"DDEBF7")
    col_colors(["45天","60天"],"FCE4D6")
    col_colors(["复核天数"],"FFE699")
    col_colors(["合计天数"],"DDEBF7")
    font1 = Font(
        name="等线",   # 字体
        size=11,         # 字体大小
        color="000000",  # 字体颜色，用16进制rgb表示
        bold=True,       # 是否加粗，True/False
        italic=False,     # 是否斜体，True/False
        strike=None,     # 是否使用删除线，True/False
        underline=None,  # 下划线, 可选'singleAccounting', 'double', 'single', 'doubleAccounting'
    )
    font2 = Font(
        name="等线",   # 字体
        size=11,         # 字体大小
        color="000000",  # 字体颜色，用16进制rgb表示
        bold=False,       # 是否加粗，True/False
        italic=False,     # 是否斜体，True/False
        strike=None,     # 是否使用删除线，True/False
        underline=None,  # 下划线, 可选'singleAccounting', 'double', 'single', 'doubleAccounting'
    )



    align=Alignment(horizontal='center',vertical='center')
    for i in range(1,sheet.max_row+1): #遍历行号
        for j in range(1,sheet.max_column+1): # 遍历当前行的所有表格
            sheet.cell(row=i,column=j).border = border  #将当前行的每一个表格填充颜色
            sheet.row_dimensions[i].height = 30

            # sheet.cell(row=i,column=j).alignment = Alignment(wrap_text=True)
            sheet.cell(row=i,column=j).alignment = align
            if i ==1:
                sheet.cell(row=i,column=j).font = font1
                sheet.row_dimensions[i].height = 85
            else:
                sheet.cell(row=i, column=j).font = font2

        i=i+1 #遍历下一行


    excel.save(r'C:\Users\wb\Desktop\LXGG1.xlsx')

# write_excel()

# import base64
# import re
# src = "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAugAAAJ8CAYAAACsrVTFAAAgAElEQVR4Xuy9B5gkV3nu/1V3dZywsyvtKqMcAElIgAISCigjZIJtjMki2CAyGC74+tq+vv7ztzG2ARtMFmDAXIwJJgiUs5CEQEII5bQKKG6Y1LHCfd7vnNNd09sznWak3um39DTTO11VfepX1cNbX7/n/bw4jmPhQgIkQAIkQAIkQAIkQAIkMBQEPAr0oTgPHAQJkAAJkAAJkAAJkAAJKAEKdF4IJEACJEACJEACJEACJDBEBCjQh+hkcCgkQAIkQAIkQAIkQAIkQIHOa4AESIAESIAESIAESIAEhogABfoQnQwOhQRIgARIgARIgARIgAQo0HkNkAAJkAAJkAAJkAAJkMAQEaBAH6KTwaGQAAmQAAmQAAmQAAmQAAU6rwESIAESIAESIAESIAESGCICFOhDdDI4FBIgARIgARIgARIgARKgQOc1QAIkQAIkQAIkQAIkQAJDRIACfYhOBodCAiRAAiRAAiRAAiRAAhTovAZIgARIgARIgARIgARIYIgIUKAP0cngUEiABEiABEiABEiABEiAAp3XAAmQAAmQAAmQAAmQAAkMEQEK9CE6GRwKCZAACZAACZAACZAACVCg8xogARIgARIgARIgARIggSEiQIE+RCeDQyEBEiABEiABEiABEiABCnReAyRAAiRAAiRAAiRAAiQwRAQo0IfoZHAoJEACJEACJEACJEACJECBzmuABEiABEiABEiABEiABIaIAAX6EJ0MDoUESIAESIAESIAESIAEKNB5DZAACZAACZAACZAACZDAEBGgQB+ik8GhkAAJkAAJkAAJkAAJkAAFOq8BEiABEiABEiABEiABEhgiAhToQ3QyOBQSIAESIAESIAESIAESoEDnNUACJEACJEACJEACJEACQ0SAAn2ITgaHQgIkQAIkQAIkQAIkQAIU6LwGSIAESIAESIAESIAESGCICFCgD9HJ4FBIgARIgARIgARIgARIgAKd1wAJkAAJkAAJkAAJkAAJDBEBCvQhOhkcCgmQAAmQAAmQAAmQAAlQoPMaIAESIAESIAESIAESIIEhIkCBPkQng0MhARIgARIgARIgARIgAQp0XgMkQAIkQAIkQAIkQAIkMEQEKNCH6GRwKCRAAiRAAiRAAiRAAiRAgc5rgARIgARIgARIgARIgASGiAAF+hCdDA6FBEiABEiABEiABEiABCjQeQ2QAAmQAAmQAAmQAAmQwBARoEAfopPBoZAACZAACZAACZAACZAABTqvARIgARIgARIgARIgARIYIgIU6EN0MjgUEiABEiABEiABEiABEqBA5zVAAiRAAiRAAiRAAiRAAkNEgAJ9iE4Gh0ICJEACJEACJEACJEACFOi8BkiABEiABEiABEiABEhgiAhQoA/RyeBQSIAESIAESIAESIAESIACndcACZAACZAACZAACZAACQwRAQr0IToZHAoJkAAJkAAJkAAJkAAJUKDzGiABEiABEiABEiABEiCBISJAgT5EJ4NDIQESIAESIAESIAESIAEKdF4DJEACJEACJEACJEACJDBEBCjQh+hkcCgkQAIkQAIkQAIkQAIkQIHOa4AESIAESIAESIAESIAEhogABfoQnQwOhQRIgARIgARIgARIgAQo0HkNkAAJkAAJkAAJkAAJkMAQEaBAH6KTwaGQAAmQAAmQAAmQAAmQAAU6rwESIAESIAESIAESIAESGCICFOhDdDI4FBIgARIgARIgARIgARKgQOc1QAIkQAIkQAIkQAIkQAJDRIACfYhOBodCAiRAAiRAAiRAAiRAAhTovAZIgARIgARIgARIYDshUKvVpFwuy9zcvMzOzsrM9JzMz89LuVyRWq0qlUpVqtWqxHEs+XxecrmsZHM5yedyMj4+JmumJmViYkKfFwpFyWYz28mRj9YwKdBH63zzaEmABEiABEiABLYjAvV6XebnSyrGt2zZquIcAlwflbrgdYj2IAj0EUWRPrB4nifpdFofvu9LNguxnlFRnslkVMAXi0VZu3ZK1qyZ1Of4PZennwAF+tN/DjgCEiABEiABEiABElACYRiq+IYQR2XcVMRrKsLdwwn0WjVYINDDMNLtFxPoEN8Q50ao+5LL5fQ5fprnGSkU8jI2hup6QX8Hcc/lqSdAgf7UM+c7kgAJkAAJkAAJkMACAqh+G+vKnFQqFSu8UR2vq2CHOMfPeq0ulYoR6/g3BDnWwfZ4jgfsLa6CnkqlGhV0CHRTSc80RXkOz5siHaIc6+EnRPrYWEEr7RTqT+0FS4H+1PLmu5EACZAACZAACZBAgwAsKqVSSYU5RLaKcLWtQJgbEQ7Bjuf1mnndCXasX6+bKnrT3iILKujb2lz8ht0lKczVr57PSj5vqunGv26e+35KxToeEPhcVp4ABfrKM+Y7kAAJkAAJkAAJkMACAsZbDgsLJnc27SuojjsR7oR5pWw957aSjvWxPcS5Eel1FeWmeu41KujtqugQ26iQuyo5LC4Q5RDkeDghjp/5gqmsm3WyKs6bop1CfSUvaQr0laTLfZMACZAACZAACZBAggAENSrmTRuLEef4N8S5mwSKdUz1vCpOoCeFvJkUGrZMDhWJQmkr0FFJh8D2M8bykskstLrAe45JohDm+Nl4PmZEOwS6EfaYdGp87BDrsNBwWX4CFOjLz5R7JAESIAESIAESIIEFBFDhhgiHAHdV72SlHIK8VIJIN+uoiNfKuRHvzvrSTG0xfnP3QHALvOdx5C2wuLgqesPq4luh7vsNkW6q5/CcG2GOB2IYMVm0aAV6segsLxDpsL34Ks7NxNOsJsZwWT4CFOjLx5J7IgESIAESIAESIIEFBCCaIaqTwhz/huCGIC+XTTUdHvRSqSql+ZKUyiVdv1I2wrzpS4e1xdhajDCPG6ktuAFwAj05ACec8ROCOpWy0Ys6WdRU0bNZWFwg0E2CCwT6xMS4PtfH+JiMjxcbE0dzObNdKuVrNR77hkh3z3kJDE6AAn1whtwDCZAACZAACZAACWxDwFXNjV/cTOR0gttVzM1PM0l0fr5sBHrJCnTrT8c2LufcRSmaOMVYENjSEOe2gu4G4tJcXBXd2VEa2egZ5KOnTexiNiPForG3QJSPj4/rwwl1NDfS6rqmuhgvuu+bVBiXFOOEOqvpg38YKNAHZ8g9kAAJkAAJkAAJkMACAhDQqIJDPDuB7hJZTMa5yTk3wnxeZmZmpDRfaYhzrOM859i+aWdxk0Fja2VJNQQ6BgCLy2ILhLN7pNOepH3TwMg1MUIF3TQvKlhxPqFdRyHSJycnbTUdPnVTcXeRjE6g46fzujOWcbAPBAX6YPy4NQmQAAmQAAmQAAk0CKBqjWo3xDVEtXuOf0N0G3GOivm8dgfFY37OdAotlaz33HrOXVpLskMo9o/KuYpx+9zGntvfJU5GLBLjv8QKTZEuWvl2Ar2ZfW4mhcKDjgo6hDl+otMofhrLCyrtJh8dlXfsA4LcWGiMSHdVelbT+/twUKD3x41bkQAJkAAJkAAJkMACAhDCTpi3q5xDgKNaPjs7p4IcVXM8n5ttVtIr8J5rQ6Ka8Zq7+MQokggWFiu29Yf+j6mYm197KsjdkhTmyeZF7nUnqJOTPV2DIghwVM6dSDcC3VTTIdDxe1hekKXuRHpSoOM9INIh3inSe/+gUKD3zoxbkAAJkAAJkAAJkMCi4hxi2HT4NBGKLpUFthYI8+npafNzqxHoWkWfn9f1IMzdpFDXFTTpMXdi1wh0aYhfI8A9iVvSVHTiaLLEnhi1q6Ybi4ovWZePjgmjSHKxE0QhytesWSOTk8byMjkBb/qEFIommjGXM7nqyap50k6D1yjSe/vAUKD3xotrkwAJkAAJkAAJkEBbce6EdNLagoQWNwl0dtZ4zbdu3SozM7MyM20sLibBxWaj2yZETpw3K+ZxQ+Q68WtSW8xQGut5qW0E+UKB3qywQ9DrvgQJL6h2Z8RHPnou02hYpBNGJ4pqdcFDhbq1voyNFTThxXQdNVYXTDrFvrSen/C8U6T39qGhQO+NF9cmARIgARIgARIggQYBZ2VxlWr82+Wcu0xzY2tB5dxUz7ds2SKzM/ONSroT59UK0lrMhNDWqrcTu+6NG5V0m3/uRLrV64tWzWOn6NUa4wQ6fOPGPw6BnsmalBYzYRSpLgWZmJxoCPQ1qKJbb/rEpBHomDiKNBh0KkUjo6Q4d2OGSGdjo+4+PBTo3XHiWiRAAiRAAiRAAiSwTeUcCStJce4Ee61mss9LJVTNjd9869bpFoE+I3Nzs7Z7aE1CFedBYxIoHOVJGwo0teeZ6rRO/8Qk0JYquhPgTWsLhHjSl263xna2ei72p6dWl5SkNSPddAqF+DaRi8WGSJ+aXNOYNDq5ZkLGxpDsYpodoYKeyWQbk0WT1hY8ZyW9uw8RBXp3nLgWCZAACZAACZAACTQIuAmh7hfO3tLsElpXX7mLUIStBY8tm6f1J7znJlqxJJWq8Z6HYd34ytUeEmtTocbES/wOCSn6amqhQFefuWdvFJrV99ZJokm7DN4lVIGvb6ZvahJYINCNXcXFKEKgI9kFVXN4z9dMjsvU1JR60rWSPmG6j6KKbvLRmw2M2lX+KdI7f5Ao0Dsz4hokQAIkQAIkQAIksECcu26eEKCuWr3Qe15rTP5siPMtW2TrlhljcZmdk/m5OSmXylIPahLUa+LZmZ9qN0FOOSwnNr7QdQIVgYml6TM3dpjmGBaKcFtht/vVpkZ20iieh/aIkNqowt/cGWiVHhM+UUWHmC4UTIXcNC6akKk1E7J27VoV6PCkQ6CbVBeTo95uwmjy8nHxjpw4uviHigKdf3BIgARIgARIgARIoAcCTpxDDDuBjgq6s7cghaVcrtoYxVkV5Fo937KlUUEvzZmOoaieB0FN4jDSujiK2bCZ6IRLVLMh1J1ITy20qjh7i7tBgL08ML4Xc9OAcryNY3TrBHHUEPSh7UQa6rrWCOMl/Oi2iVGuOKaVcWd1WWtTXRC9iEr6mqkJm5Fe0Eq78aI3s9HbCXEXwdgD9pFalQJ9pE43D5YESIAESIAESGAQAq5KntyHaRhkJociVtEJdEwIhY1l8+bNKtD155YZmZmeltJ8WcqVktQrFYmiUGKJJOsmavoprUL7KU8r6BDsxuoCm4vLPTciPLKTRNViE4mEKsBNMLoR5UZ5NxJmYtxIxBJgvHpTgao67C5mXdOItNlsSG8UckhpyZkmRWNFmRofk6kp40NHJX1qrbG+wKcOgY4quhPoEOKLVco5aZQV9EE+i9yWBEiABEiABEiABFTwQoC3LkmBjuQWCHREKrrM802bNqlAdx70Oe0aWpJauSIhqudxpJXzXBrCPK2iPOv7kkHkoe8EOqwnLvfciW9P4iiWyIpuCHTXvwi/0wK6vXlQ8Y2mR1FoOpxGsdTtTzwPVeh7orcKtlCvnvQUxpExXUMLRZkYK8qaiTH1nq9dO9WooMPqMjEBAW9SXfQGw/cbHvrFLh9U22l12ZYOK+j8g0MCJEACJEACJEACHQhA6LrEltZVk/YWCHQ85uZKDVHeqKBv2qrVc0wcrcJ7Xq1JGNUlLbFaWbKZlOQyGclBmCPqEI2DfAh2VM8xYVQTyxtvb8S3qd43fPD63HjN8QSvx/Z1UzkPpR4GEoSxBEEk1cCI9SCMtIoehiJBI1fd04mpqILncgXJWpvL5MSYVsydQF+7DtV0I9Bdh1EIdFchXypakcku7S88CnT+SSIBEiABEiABEiCBDgSc77zdalqZDkMV8M3GRKWG9xwCXX3o1t6C+EUI9AipLWEofko0fzyf8SSfzag4z+eyKtRRTc/4mCxqbC7iqVNdh9GMd2xaWNQuo9rcCPVmBd10N9WUGXQ5rWO8kdRCPEKpQ6iHxvYSwPYCYa/GGxe7aHLOkdYCca6TRafWqMVl3Q5TttMo4hjHtYruBLrzzy+Fl350VtD5B4gESIAESIAESIAEeiKwmLXF7STpP3cCHRYXNzlUBfrmLTK9dUY96bC21GoVicJQvCiUbCYt2YwvxXxa8hDn2YzkkUOeQaY4rCJGoKchznUSp41HFM8I6YaNBc+dQHc2GFNh15uIIJR6BJEeSbUeSK0eSKUeGpFeN+K9FliRHoZqdYHtBW+YyeXFz2StD71ZQYdARwXdxC6alBcIdJcA041AB0daXRZekqyg9/QR5cokQAIkQAIkQAKjRgCVcQjcxZamQK9KqVTWxkPaNXQLklumdXIoBPqW6a0yPzcjtUpZInjZ40DSnkheM8fh885IIZeVQi4nxSxEO6rnGfGtLx2pLurXxsPMA9XQRZjNIdSN5xwuclNRdxaXpkAP1OICMY5HtV6Xas09N6K9GtStiMe6xpuO2anpTE589aIXZXzSWFnWYYLo1JRW0U3sohHoqLIn4xa76R6KdXAzwsUQoEDnlUACJEACJEACJEACixBw0YlLAWrmnxuBjgmgEOhbNyNe0aS4mOxzdA6dkXq1LHGtjvq3ZFKeFHIZKWSzUsjbB0R6FtV060G3kYVIczHdRU0UolbObVoLbC+t1pZmZd1U0AO1saByXpd6PZQKKui1QKoQ7LW6VtNVtEOkB/CpR1JXq4unAj2dzkq+mJcx7Sw6bhJcrEDXuMU1xvrSj0AHX6a6NK8yCnT+SSIBEiABEiABEiCBNgSWmhjqVm+NWMQE0KRAhzDfvNkkuMzOTkt5fk6CWlXioCZpz5NsKpYiBDmEL4R6PiPFXE7FuXrQbeQiKsywi6Q9k73uSVoik3Ruauj2ian0NyMWkfICP3poLS4Q6HXYWzQSMpAyquYQ6TWI9ZoK9kpQV6FeDyKph5FEmJqa9iXtZyWXz0th3FhcIMrdw1TQp6zFZWHUYjcVdK0aw0qTyTDVhRV0/j0iARIgARIgARIggfYE2mWet665mECHIJ+2XUNRQYf3fGZmWiqleQnrNZGwLtl0ygr0nBTyORkvZqWYhUg3At1409PNpj/ISddOn+YBe3hr59AF+ecQ6jblxdlwkJMOy46xuFhhjn9XAylXq1IJQk12qSLPHRNJ4UnHpNRUWtIpX7KFvOSLpoKuVXMbtZgU6K6jaLJZUbfXmNum2/VX63qsoK/WM8vjIgESIAESIAES6JtAN9Vzt3Mn5N0EUVTRIcjhQYfFBVV0NC2an52WahkCvS5eFIjvRVLI+Oo5nygWZCyP5xnJ2wq68aCbFBWtYKNZkfWhm+xwE69oKujJLqMLnzsPuhPpZjKoTZ2pmcmi1VpdqtWalJBEA/Feq6mIL9cRv+hJrBV8XzI2D318YsJkoU/t0BDpqKCbmEXT2KibmMXWE8QquiFCgd73R5cbkgAJkAAJkAAJrFYC3VTPWwW66SBa1pxzCHR40CHMIdDxb9hbKuV5iep1SUWBVs8xQbSY92WskJdiPifFnC+FbE5yWTQIMtXzVDptBHo6rW+paS66RCZ60Qp0PGsV7NvELNqOp0akw4sOa4v5iQo6qubwpkOsazW9FkkNOeoexpDRNJdsoSDjE2tkbNLaXNasU5E+MbFGBXqhYMQ5kllMB9Rmdns31wur6BTo3VwnXIcESIAESIAESGCECPRSPQcWJ+adQIcHPSnQIdKNQJ+VaqUkcWAEOjqH5v2UFHJpGc/nVaTncxkV5/CfZ2zlXAW69aBrd0/1oOM/I9LdkpTBTqgnmxhpBd1mtmseup0MaiwvtopubS9lWFzUjx5L1US5SJxKSyZbkMyYSXGBIJ+YWiNrJtfK5BQaFa2R4tiYFAsmM91V0HvtFIr1sf0oL6ygj/LZ57GTAAmQAAmQAAlsQ6Cb5JbkRs0Ul5pW0CHQZ2dNzKKK81lMEJ2V6vy81CtlCTFBNI4kl/Yk53ua2DKmvvOcEejaSTRjmxOhcp0WL2UesLggmlEj0T30FTVWF7d4LVYX18wIPzFO0y00TDQtCkzzogBpLjWp1SJ9XkZH1HpkLS4iITqYemlJw7qSH5OiTXKZmForkxNTKtDHxycbAt3kt/t6Y9GrQMexjHqiCwU6/zCRAAmQAAmQAAmQQILAUl1D24Fy3u5KrWo6ic5XZG5uTmYxMXTrFpmd2SKl2XkpIcGlWpE4rEsqrkvGi6WQSWsGOpJcIM7RsChrxS1iFf2Ub33nxipiRDqiyU0eugp0JLvYPJeU/akV9NhEMZo4xkgnjAZxYJoWWZGuNxchRDqq6EhuQS46/OjIRIc/Hc2NRGox2p1mxUtn1OJSGJuQMTQlmpwyAn1y0vjPx8YkX8irOHeJLP0I9FHvLkqBzj9JJEACJEACJEACJGAJdOoauphAh9BFBTop0OdmZ1Wgl2amZQ72llJJ6tWKCCwusUlxyaVFMlnfCHTYWqzv3FSfYW0x9hatlmv1PG2euwq69XcjUx1LylXTMbFTGxclO416EuK/ECI9bAh1vcEI641KusYwai56oDGLtciTmia5ZMXzM5LJImpxXIrjkzI+MSVjsLpMTKhARxfRTM7YW1wFvd+La5S7i1Kg93vVcDsSIAESIAESIIFVR6CXyaHu4F1F2gn0chlV9LLMzc7I/Awe01Kan5VKuawWFwlrkopDyaQiyfkpyaZRRTcCHbnnreK8KdAhzq1YT3nG6mK7iqKgrq1ETWr5AnFuRDq6iyI33dhbkI3uRLobPyrpeA0WF1TRteOoFeiBetBNFno6V5D8WFEKY5MyNj4pYxMuvaW4QKCjCt7rBNHkBTXKNhcK9FX3p4UHRAIkQAIkQAIk0C8BTJg0zX66X5y/W6MJq1UpVcpSmi9JeW5eoxXnZzFBdN4I9GpZPAj0KNSYxYzvaTdR+M9N5nmzam6EuZ0Uip+NSjoq6Pi9CHJdzKRR+9ND9GKzcm46jMZGnFurC4S5+V3zp4litN50iHRtZoROopHURSSIPPFSWZT7JZ1DFnpRCoVJKU5MSGF8QsbHJqUA60uhINm8SXGBQO/H3uLI4xixn1FcKNBH8azzmEmABEiABEiABLYh0I+9RevWKnZNAyBNcqmUpVwqy/zsnJTmZkz1fH5eE1zq1ap4QVW8OJCMROJ7nk4Ghe9cM8+16ty0tUCFN0W6yW3R6EK1uaQkjTQXL040MGrmopvKuRPoEOCRRAlhbl7H7yDqzetOuGsVPQgFAS71WCT00hJ7vqS0gp6TXH5M8mPj6kUvTEzKWHFCxXk+n29YXPqdIJo8MaNqc6FA5x8oEiABEiABEiABEoA5xIrsXmEkBXqtZnzoeMzPzUp5rpl/Xi2VJUAFPUCjorr4qVjSEouP+ELfFx9V8TQmg7rKs5kIqkuigygmiEKgm8ZFsLrguVkPYt1ELHpWcDdvIKIwRmujRuXcCXRzk4FXzLZxHEqARkahjWWEOIe7HePys+JnC5IrjEmuCJE+oSJdK+r5guTyOcnmzCTR5RDoo2pzoUDv9VPI9UmABEiABEiABFYlgV7TWxwEJ9C1+Q/yw6smzaVcmpPKHKrnc1Iul6RWnpegVpG4XhUvCiWNCnocq5CFQIfQVhuLTvw0uSymHm6EOrS62lmcKIeYRwVeJ49iW21BqRup5xxiu1EVR7xiawU9YeXBvlWi40YFE0gRyQhrS2jEuedLjOq+n1MPei5flExhTMU5xDo86aiea/55dvkE+qimuVCgr8o/MTwoEiABEiABEiCBXgn04z831WdjcQk1ohBZ4jWpQqCX56UyNyPVEvzn81KrlCSoVo1AD+uSjmPx4shUwdNpFdlGhSMs0VhToM1NPRyxLSb/HGkumuiSEq2iG4EOYQ//C9S52XZhF1E7Rp0g6irlEOj2FsAKf/w7iiHKjcAPtbKO6nlGBGPM5LWCnoEYL4w3Kum5YlFyOVTP4T/PaQXd3Uz0eh6S64+qD50CfZCrhtuSAAmQAAmQAAmsCgL9+s+dQEelGgkoTqDXKjWplpF9Pqs/EbEYVEua4hLVa5rkkhToLu3ETKq0EYm2oq0SOkbMoZa51beO9WGJ0az0pED3bNfPht8c3nJj39Hsczsx1OSkmwx1iHIT44gqvBXv7r1RiccNgwr0ZopLplCQTK4o+eK4ZNXuUpRsNicZ20FULS7aTcldHo0nPV8vo+hDp0Dv+TLhBiRAAiRAAiRAAquNwCACHSyiKNAs8Xot0MmidXjR50umcl4yP+vlkoS1ioQ1VNCR5IJIRGMz0RwWVL+dOEdVXqvz5oHXTUUaOt34z7XqnkYKDDzrxiqTUsFtGhSFVpQnBToaFcG6IpFZz1jcsS0EOvYNsW4nmup9QVrL9qigp9K+etAxSdTPFySbgzCf0MZF2fyYCnRjccG4jI8eD3wbkFDqPV86FOg9I+MGJEACJEACJEACJLD9E+h3gqg7cgj0MIIwR0ShEeiomtfKJaniUZqXehVV9IpE8KGHdfGC0Aj0GFIZCeaRFrXV3hLBauLiEq2Et1VuFdMpk/6iFhD9t3luBDomiFpLi8Ynmiq6dgyNAqPNjQY3Nwe2MykEOmwzbl6qsdukRVIQ6RkRVMVtzGImB5sLJotOqFjP54viZ/Mai7jcAn0UJ4qygr79/03hEZAACZAACZAACQxIoN8Jou5tjcWlpvnhbrJorYTqeVkqFVhc5iSqVjQHHVV0FehhKBJH4qlihkA3nnMV5g2BbrzgWkO31XMnxv20qaT7CYGO5kW6D2draVTRMb5YUEE3mehWnKuxPbZ+dvseNhFGbTe2gi4pVM8zkspmJQ0hjomghaL60LO5omQxaTSblWw2L6mMqe4vVwV9FCeKUqAP+IHm5iRAAiRAAiRAAts/gYEFOnLEQ2Nvwb7qlapaWtR/XjGVdEQswoMOge6FgcQhxHKQEOgobLsJnFaoW4uLqXQby4jmn8N37iPBBXM30yZyMQV/OhJXjOfcVc6d1UW7pDbyzk1VPqXmmVitMlooTyTFuPeCzcVDNrufFy/ja6OibLagAj2THzfiHPnntoKewjoU6Hqc7lUAACAASURBVAN9KCjQB8LHjUmABEiABEiABFYDASSvmPzw3hcz4RIxhkac66MCMV5RewviFfGoVysq0rWCXq9JHAWinYAwsTNGcgsq6p56xDUiUSvpxvLiJls60YzqedLeookuEOp2RYzJiHHRyaEQ6wEmsYaBqdZrt9SUWmvUe44qvBXo2gTJRT4a/4taXFBB97M5SWlSS178/JhkbNxiNlcQP5uRbAYVdN/muZsuqIN60EcxyYUCvffPIbcgARIgARIgARJYZQSWS6CHQSBhPZBatSxBuSK1ypz6z2vlstRrZQlVoFcltALdQ6Rh4sbApavgVyFEup00CuHuJokiHMXXWEZXPbfdRbWyDlVvJoniUddM86ZQD2P40E14SzNgxUwSNY2PMCfUWV0g3/FvdBFFBT0jaZ0kmhc/k9ccdD9n8tAz+B1iFrNZ7TaKhkuNCvyAk0SxH+x3lBYK9FE62zxWEiABEiABEiCBtgRWRKBXygIfeh02F7W4wINe0pjFKEDUYt2GGZoh2fAUiWOvUUHXCZ/ae8g2KtJJnDZmUYvbJg/d+dJVoGMOaGSy0J2lpSHS1YNuPO76ni4ZRvPVm+K8+ZppnqSNklIZSWOSaLYo6RzSWsaMSEclPVfQ30FIpzM5CvQBP2cU6AMC5OYkQAIkQAIkQALbPwF0/+x3SVpcXAUdXvPAes9hb4FAD+0k0ahelTCoi0Q1nRrqElVSmJCp0Yq2i6j1n5uGRZ7EiW6i8IzDzqJdRa1oh/8cOeZqK9EuogtFuqnK20ZFMKrbBUkyLhnGxrAb8W7aI0lslLtgoigEup+BGM9p9RwZ6BDofg4edPwOAj3fEOia625anPaLV7dDE6RRWijQR+ls81hJgARIgARIgATaEhhEoGOH6ukO7ATRel2iWtV4z1FBr5ioRfjS1YOujYrqEkd1kTjU8cDzHUEOq73EdhFF5mIbYesq6CazvDlxFPnlptFRc0l2E8UYVfxv47XHdrYy37C7m3x2o9RhccHY0hqzqAI9mxU/V5BsAdV02FzQuCiv1fNMtqB5jTq2lOmKapPe+776KND7RscNSYAESIAESIAESGD7JLAcAh3RifCfI8nFCHQzQTQom0ZFYaUiQa0sQb2qYh4CPeUaFUGUo+EQ7CRuWUKgm86fxqYCe4sR6rFWvFtFuhHmNh3GVs6dSNemRFacqxa3+t68ju3wOyTDwE/jq78cFXIfQh3NigpFU0HPY5Ko8aHDn06BPtjngBX0wfhxaxIgARIgARIggVVAYFkEOpoAWYEe1CoSlMvbCHR40EP4z1WgI1ElNg1E1UjerIYDKardyYZCDrPWo60gdxV0tbvojppxjE6Em8QW0/TITSBtnDKtupt/NYU9xDzEuUl6wdIQ6OmMxi1m1OJimhWZyaKwvRQ0Cx056RTog30oKNAH48etSYAESIAESIAEVgGBQQU6GhVpN1FELNbqWiUPXQdRl4OuFpeKBEhwCWqCBBf1oNsqeKMbaKOKvq3FJWlP0eo3tL31orcKdHda3DZOoDux7tZvFecqyBNVdyPQU2pzSfs587AWF8QsahU9NybpbE6Feyqbo0Af8DNBgT4gQG5OAiRAAiRAAiSw/RMYVKBrSkpYlygMjc0FFXQIcmtzqZYxabQsqKwjZjEKa+JFITJaFgj0hr9cRToq6qaC7SwqTdIQ57b67qrmiQq6E93YboHlJcaM0qbn3OzPJr9Yb7qJdTeNjoy49ySyPvRUGkkumBBq4hY1YjEPgW4sLhDuEOqsoA/2maBAH4wftyYBEiABEiABElgFBAaJWXQVZ1TQIdBRQQ9RQUfn0FLJThI1Ah3CXQU6KujoIprwoKPRkPOTm5/N9JNtBHoMv/nCxkqa4KKecWOVSS7Jfydfb1lN4x1Nk6RWgY4KOjzoyEKHvQV+c2Nx8SHQ85g4mhMfAt3PGoFuvfGDThJlDvoq+IDxEEiABEiABEiABEigVwLLIdDhKdeOnUGgk0RNN9GSVtFdBd0I9Ip61U3MoslpgZhtdPTEcyvQIdKT4tzYVFAER+W9Wfk2vvSmB91V0N3PhaLcNT3CVgtFvk4NhUhXy46poGPbEDcL2qwoK55vrCzqObcedFTT4T33MxnxnUA30TADp7hQoPd6NXN9EiABEiABEiABElgFBDR5xU6m7OdwjKANBF70EDaXWlXiWkU0D71FoMPmEtXrEoc1eEmMSNdUFiSymCZESwv0GCq64V9XEd2lQHeV9Gbyy8KjdRV0J9Dt4LSTqMYsQqBDiKNaDnHuJofmipKCLz2TkXQ607C4LIdAx1gzmUw/p2W73YYWl+321HHgJEACJEACJEACy0UAVW8I674XVJ1DiHQIb4j0ugQV0zk0qFSlVikbywty0KtVnSgaBbgpCCQNL7mtNqcQtYjJmKig28mfGJNrhoTcdFNRX1hBN0La+NmdzQXP02heZC0vjSq69g5q/h7qXh3tLorRVtDd+6JBUuT54qn/PKuTQP2szT2H9zxXVHtLOptRIR37GfE8HMe2Vpt++ML64/t+P5tut9tQoG+3p44DJwESIAESIAESWC4CEOcQ6X0vDVuI8aHHYaAiXG0u1apaXcJKWWrIQq8bH3oMK0xY15hFY09B9dxMDNVKOgSuwCJiMsyNtcXYT5o55fY1K9AbFXLb4Egr81bou6o5qtqerdRrUyTddluBbqahImIGjYrS6i33sllJYZIoPOiopOeLGruo3vMs7C2+xOmc7n+5BDr2CZE+SgsF+iidbR4rCZAACZAACZBAWwIQvbC59L2g+hxhgqXxocfoKlqvSVitSb1akXoNsYvzUtPnRqALqvZBXeMWjac82XTIZqLbrqKugg4d34xANNX05hI1Go86Ud7oNGor6VqpT+SttxPo2J9W061A99K+RCkI9JxW0NNZk3musYpZUz1P5bIqzlMQ0qls4z1aJ6v2wxdVedxcjNJCgT5KZ5vHSgIkQAIkQAIk0JYABCkmiva9NKIJAxW3GrlYRyZ6TcJazcQrVkqCOEcIdlTRY2tzgZiXONKGQaZm3kxiUUENER4ZwYznJnM9khgJMNaWYsZtRD52pNX3xKORDoO9twh04xO3nUYT4tzeMYikMuL5aFCU1XhFL1vUCno6g4mi1t6SgTjHA/tvVtD75pnYMIv3bI2bWY4dD/E+KNCH+ORwaCRAAiRAAiRAAk8NgcEFuqk3R5GpasODHgehNi4yAt3ELqKCjmZFqKIj6QUNi0I0LkIFXavhiE9sinSbht4Q4qjSq0DXTp+helAa/vSEQDeTTj1BM9JkAyTsGwK9OUm0xeJiq+ZquEHVHZVreMohxn2TcQ5R7mfgQzfJLZgciomhqJ5jv3HKV4uLuVdYGPfYz9nM5XL9bLZdb0OBvl2fPg6eBEiABEiABEhguQgMluQCpWy84ZqAgudhIBFEugr0ugQ1M0kUdheIdAj2yHYVVd+6Cu5ImunnZsInRLUT4dDxzQjEwJjENdHFpbpoGb7p/4ZQTyPC0SbD6M6TAh3xiTbrUfeCZHa3j5TA3uL58J3nxMtkJJMtSAr+c+tDx+9RWU/7mERqs9t1ouvyTBAdxQQXvbGJF5qXlusa535IgARIgARIgARIYLsiMFiSS0KgW9uJVtPVZx5IWK+r77xWLdsKuklzieuopNfUix5HxurSiE1EM6LYVNPVE67Vc+tBx/MYAt1GLmr1HTaZROXaWVycQIeAhgWlxfqCKrurdqs4T6V1HYGoh/DW5Ja8+PCga6yiEeioomPSaBoVdls9V684IhmXqXo+ihNEKdC3qz8bHCwJkAAJkAAJkMBKEhhsoqipYru6p7G5hBJbkR4gZ71ea04SReyi2lxMRR2vwbcu1iID4Y3EFo1MNBmI1ttuBDqq5piQ6qloV5+LxiyaYrgV4bCZeJ6JbHTZ6hnTfzRpcbHmdxX3kSa8+KYCDtGt9hZYWRCjiEhFPC9oDnoKwh2vo7toOlE1txGLy3GuRtF/ToG+HFcO90ECJEACJEACJLAqCAzmQ3cCHSggpqGZI+MX14x1CPS6Vs9r2mW0pAJdrS4Q6PCoB1WJQjQ7QqqLxrU0BLruyya4NAR6GKpA1wmk2qwID2NPUQu4rZSnfFMRhwUlrc8TOeg2D93OTDWVcwh0CG5M+oSFRcV5XivnGfWdQ6Bn1ZeuAj1l0ltMlDusLcsXiTiK/nMK9FXx54QHQQIkQAIkQAIksFwEBvOhm1E0q+ihVrtRSVfhrTaXmtSDqtTKJfWgh1UzWRTJLvp6hImjdfFQSbeV8kb1HFVzgbfdVNAFlvUQIt5W0DVP3VlcTPMjneQJtZ72BEI9OWFUf++84gk7DCwuqIrrI4dKeV4r59o9FI2JMpgYikmjObOOrbi7KMTlmBgKjqPqP6dAX65PM/dDAiRAAiRAAiSwKggMZnNZiABdQhuTRmF3CWoCq0tYd51FK7aSjop6WTuLahUdDZM0etGIdPNoNijSf0Oko/GpraDbkn0zphFWcCvAtbqt1vCUqXTjpxXuKuBdpR3P4VNHRRyZ5tmcCMR5HgIdk0ORe16QtPrPYW/JimASqbPEJP3vy3A1jGL+ucPGSaLLcAFxFyRAAiRAAiRAAquDwGA2l4UMkMqikzs1Fx0C3TYvqtdNFnqtJAE6i2LiqO0wCi+6ThgNauJZD3ujMVFohLmrrCP0xVhcbEXdvj1q6Ckv1phEjVSErQUVdCfArUCHUDddRZG+YtJeYG2B6NZkFo1ULIiHynkejYnykkUFHdGK6CqKCaQaqWg968sQqZgkOKr+c1bQV8ffEh4FCZAACZAACZDAMhIYLM2lORCNTbSWFwh0CU2iS71uGhepSK9ApNtkF8QuItEFIj00nUYR1agVc9hYwlg8J9BtqqKzu8Dmoj50K5IxFdRM9LRVcSvQ44Q4XyDQIeS1um68565jqAdhnikYgY4Keq6o4tzPZETSxt7iJogu4ynQfSLBZVQXVtBH9czzuEmABEiABEiABNoSWD6bi+ns2aiA181kUbW51JDoMm9iF1E9ryAjvSJRvaICHUIdsYsestThYceEUVTLAyvWtXLetLioD91V0NWL3hToMTqHWoGulha1upifjchFtbakTFUcAtxOAvXzRUlhUmi+qJVzdBBV/zkaE7kJom6i6TJeT6Nsb2EFfRkvJO6KBEiABEiABEhgdRBAJRqTRQdvFWMEuvOho7MoElo0xUWr6GWpaQXdCHSIdZ0wWsOE0Zr60F2qi9pYMCE0tFGOSwh0vGeygu4EesPKAsuL9acbzzkertpu7S2afZ6TVG5MO4dm8mPGi66TQ439JZVC91DjWdf2RstkccF+INCXa3/b41XJCvr2eNY4ZhIgARIgARIggRUlgOQVWF0GW1oEOlwukWlYVMeEUUQulo3FpVatSB3P68hHN7GLEOjG4mJiF2F1gUB3sYoSeY2sdPwOTUIbVXRbQTexiaaCvk21POVJBKGO121WuvrPM1m1uGilPDemFfRMYcwmueS1MZH60+FDx4RS52HXwPbBl1FtTpQkR4E++HXEPZAACZAACZAACawyAstTRW8KdK2iB7EK9Hq9KgHiFtWDXpZ6eV6q+KnNi1BdR/RiRQW6hyQY+Nc1SlG0io6fKdO41MSXa6Mi8xyLEeqRVsVjF6PoJona+EKIclTWI017SelzE8kID3pG0pp/juxzTBItaAUd/nM/lzXNivyMdhaFzcUIdJu9PuB1wOq5AUiBPuCFxM1JgARIgARIgARWJ4HlqaI3LS7oFIqqvD5QMUclvVyyAr0k9Qp+Z3zocR156DURbUYUGvFtRbhOBo2MEMdz/NSOo/Zn82zERnRDPNs4xTgVq85HwAtEeqSV9rTEXkofghQXeMsh0pHWkkXEIuwtRfWfZ3IFyeSykrFNitTugo6jKtLRKAnDTNbye7s2WD2nQO/tiuHaJEACJEACJEACI0VguSIX3SRRCHR42yHQ65Wqes5RPYdIr0Ko49+2eq7i3FbQtaOo6mlvQaUcohyV9EblHM8TzYr0H1qOtfYWiOe0EehG1ZuiPF6P4Fq3Aj1GtjkSWjK2g2i2qAI9m4c4L0gmizx0TCTNiY91rEC39wIDCfRRjlZMfrhYQR+pPzU8WBIgARIgARIggV4ILEcVHQJdH1agQ6RDoFcrqJ6XpFqaU3sLfOjIRkcWumj1vC5eDP85BLqnkybTCZGeEk8r6Gl91Sz4XeMfWnK3r+gPvGZtN14ssZbcIdJ1iqf+jL20RKiip3wRHwIdD1hbipItGIGeyxfU6pJBdV0TXzKS9tD8yNxE9FtBZ/W8eWVSoPfyKeW6JEACJEACJEACI0eghsjDRIxhrwCcQFdhXq1pFR3pLSrQS2WplGebaS6ooNfRpKguEtUlFUcmcMVW0CHQNd/cinJXMddKuopjq891vViTGHVxCr7xm0jXhcUFL0aRFemeJ4GWwiHQs5LSKnqzgp7NF1WgZ/IFyeZyWkV3kYg+/Ou6v94tLhg7qudc7OmKB7niSJEESIAESIAESIAEVjmBfnPRXRdR/HSV+GoVE0TrUq7MS62EB+wtc0ag62tlrZyjqVEKjY6sQLchLGJEsKmkq1B3LhatsFs7S+N8JLLRrV8dAhq/jcV0OdUKeuxJFHvogyRorRSl0hLBk55yAh1NigqSK6CKPibZfF5ysLzkCpLN58RPZyXte42xtYtH1JuHJW5yRj33vPUjxAr6Kv+jwsMjARIgARIgARIYnEA/3UUh7LHgp5scCoGOinwFlfPSrBHp5ZIK9LAGi0tFJDSTQ9OCyaGxVtAxCdNPJYR5QqRrgdyJdfWbGy+6dZgvOHgIc30lDq3/3KTBqDjHTwTFeCnjSU9ntWlRKoOJosbikrUiPa8VdPuAH933EQCjVpde88tHvWtou6uTAn3wzyz3QAIkQAIkQAIksMoJdIpdbK0Ou3+76jkq6BDmrQLdeNBNR9Gwii6iFROtGAXiw38usaRTngr0VMqTjKucW2+5yyBXb7oV6u5UaHZ6omptnkfGew6ri1bPka++UKAHmDQapyRKZYxAhxfdThJFBT2nD+NFz6LLaC4r2bSvVXR44J1ATwr1xUQ7YxXbf3Ao0Ff5HxQeHgmQAAmQAAmQwPIQWMrq4uwspjptrCXud87e4gQ6RDoq6FWX4IJJohDotaoV6KFIFEjaC7UjKIQ5rC3mp3Gfa6KLyzR3FXTP/L5hN3fiHLGMFkFoK+hOoOs4I5EgirVZKVozoYKOW4PI81Wge+mkQEeaixHoSHXJJQU60lwwv9S+W6tQbyfSaW2hQF+eTyf3QgIkQAIkQAIkMLIEnFXFAICvuolCbSMmt7AhziHqnUCv1wKtoFcqFZ0gWinNS70Ee8u85qKH9bJEQdXknseBeBIKvOfO4pJOp7SaDqELEaxiPFE5NwkuRqCbCaNRwqNuHC/qPTfTQtWzrjX1UNeUMIoliCHSTexihCQXxC36eY1ZNFGLefGzBcmjit5icUmnPfH8tDZO0jFisqmh0QyTsRYc/MJP++IjMYbLNgRYQedFQQIkQAIkQAIkQAI9EKjWahJGmMBpJlhqFdraSZLVc5PeAjt5aLLPNb2lqiIdAl0r6JgkijSXWlXieknioCpxFEraCnRf/ec2YjGdljQCVqww13jFhOVFYxdVo9vJo24GKSrt9kYCFXTT1KhpcUEjUvjP8RMV9Fh8435P+xKnM+pFxyTRjDYsyks2V5RcDj9hcclrigs86B4sLqmUiX20Ar15A2Gq+O7faT8t2QxTWxa77CjQe/hAclUSIAESIAESIAESgAivVGsShtGiwtxEF5r0FvOIVKA7D3qtgkmi81JH9bxSlqBekxgRi0EZJW1JxUhxicVPmeozRDoq6Bq5qFaWlOk11KNAjzzT9AilflS2UfE31XMIdDOFFNNTNcUlnZEolRUvk5V0tqCdRf1cToU6hDliEdFRNJ3NaCU8hRsIVNE9Y8fRpBk7RvwbTHAsOI5cNtvzZNJRuvIo0EfpbPNYSYAESIAESIAEloUARPrcfFnCECI3UnGLZHH8HsJcf8I6EhlxHgaBVtHhQ69pHnpVauWy1KvzElZgb6lKFFQkCqviacRiIOnYWFycQIdjxFlcNGoxMSFTJ4uiou8EfDPKRY/XVdDjlKn2YwKpmmBUoJvqOZJcYG0JtKtoWuJUWuJ0XrxsTrxMXlKZrLG4ZFA1hzDPSgaNinxMEEWKCyay4iaiOanVRUI6Ow4SW4qFHMV5h6uQAn1ZPqbcCQmQAAmQAAmQwKgRgIVlemZegiBUUa5CF7aWOJJY7S3wosODbnzo9aAudQj1Wk3qtZrmnodVY2/RCaKonodVzUFHBjoeiFpUwauiFwJ9aYuLlthtR1HPinH8G2JebTgpm0eOMeLfamnx1HuueeiIWPR8iVPIQU+L+HmJ/Zx42lU0K74Kc4h0CHNUzjMCuwqEtwd7Sxoxi0aoG5uLZ24qxBPfT8vkWFF/x2VpAhTovEJIgARIgARIgARIoE8CEOdPbJ5WkQ65C8+5VtPD0P7E7yDWkYUOL3rdZKLXa1o1D+Fnr5UkqKGCDnFuuoh66CIKcQ6hDoGtPvRYfLWQWE96Gw96q0B3h+Uq6JGLVnSeea2eYzoqBHrKZKBDmEPdp7MS+xnNQYdA9zI5FehIdtH0FZ3kiZ9GlHtIcIFIXyDQTUU946dl3ZoJFfBcOhOgQO/MiGuQAAmQAAmQAAmQwKIEILwfeXyTVOuBqaRDkKv1JWF3gUBHJR2WlyAwD9ha8NDqeU0ieNDV4lKXVGREetq0DLLe86ZAT07GdIJdU1wSKS/wvLhJmRDoOp6EQMeNBDzoxtpik1u0gygEui/iQ6CjWVFOvCxEOqrlyEXPasU8pdGPvlbONbEFb5+G99w3XnOb357N+LJhhykmtvTwGaJA7wEWVyUBEiABEiABEiCBdgRgY9n48GNSKlXV4gKBjgViXSdiQqCrYDfVdAkh0jExFA8r0AMj2GFx8aKapPBICPSUh6ZFdpKoTUtpWkkwadRmMiYsLhiDCnjri9c6v62ehxiL9Z9HMW4FPCPOEa/ooy1oTkW6oGoOD7p2FjWVc4hzrZqj0m4z2vV98E8PrxlbznghJ7vstIOkbeQir57uCFCgd8eJa5EACZAACZAACZDAkgQgeO978DHZOj0nsQpi40OHy1ttL/qw3u+wriI9CusSBYEK9TAsS1zH72v6QBXdpLlEkrZZ4mmkukCcp1LiJxNcdHKm8X2bqHHz0y2YQKo3DhDoqPLbsSD7HCYazT1H9Rw+FT8jnuafZ0WyebW6pDI5SeF3+jDiXDPXsT7eq+GNd2NIyeRYQZ6xy3p6zvv43FCg9wGNm5AACZAACZAACZBAOwKoTj/wuyfk4Uc3aYfOBR1Gnb0EdesolBgCPQi0a2hYDyQOEbOICnpgvOgBBHpd0ugqagW6WsO1Oi2SsX50V0VX20mii6c2K2rpagqfPKr9KtZj05gI3vNAFTbsLaieZ7RyjtSWCJND/ZykrGiH31xSGdsQyW6jY3NNk4zdZacdpmTX9WuZ1tLnx4QCvU9w3IwESIAESIAESIAEFiOweeus3HbXgzoxtFnGNp07Ne9FbTCheFEgEqCSHqr/XBsVBaa67mkVPTCVdO38aTLRXXfRjHq8XaShjThMTBzF+xqjja2aq6XFZLNrRjtyz5HcAv85POcqvn2b2uKrrQUTQ9V7ns6qlcWzAt0knKPinjg8dAf1M7Lv7jvLmokxXhwDEKBAHwCe27RWq8v9D/xOdt91JykW88uwx+Hcxe8eeVwefuRx2W2XDbLrLhue9kGiSUSlUpV8Pif5HLuRPe0nhAMgARIgARJYQKBSqcmtdz8g0zMl83tUs7XTkPGBo4qujzCUGFV0eNL1URcJaiLwoEOoq9UlVIGOzY1Aj7WSjiq6Zo9rBmNa1HDSqGY3K/jO+66TQtE4CeJcrS4ioedLhH2gcyiq4xlMBs2JWIGO6jkmhqrnHPGLamsxk0LNDYB6amSiOCb777mrYFIol8EIUKAPxk+3vuXWu+Tv/vGLMjk5Ln/54bcPhXhdhsPaZhfnXXCFfOXr35c3vf4VcuZpx6/EW/S0z2EbT0+D58okQAIkQAIjQQBC/KFHN8m9Dzxm7CzIIdepmTqD1D5gd4FIhx+9pr50UauLTXPBvzUXPVJhjsmXvofn0OR4bnPHMXlzgd3ECXR44c0EVcxdRZJMBJuLiASIWNR8RF87h6Jarj5zF6mYwcTQrPrPzQxQCHNfBblbMGF0953Xyy470tKyXBc1BfqAJHEX+oWvfEcuufw6OemEo+RP3/RKjR5ajctyC+JSqSKlclkeevgx2bxlWu7f+LDce/9DMrVmQt7xJ6/u+G3Eco9nNZ4zHhMJkAAJkMBwEKjW6nL3fQ/L45tntIKORSeSRrC5mMZG8KTHmDQa1iTGRNEAmeg1Y4PBepqJjsZFIplkNd1W0JOJLrp/dAv1vEZyjKucm5QZpMp4Emvuua2cIzJRE1uy4mcRp2jsLfCkpzwkt8Cjjmq5qZCjGr9+3Rp5xi4bJJvJDAfoVTIKCvQBT+Q99z4of/uxz+rXRe99x+tlv333HHCPJuR/fAg7bQ0iiHHn/s1v/0huuPG38uSmLQJb0GLLWLEgf/nhc2Tfffawk2vAZNuuY4OMZ+CTxB2QAAmQAAmQQB8ENk/Pyl33PyylctXaXFDSDsxzxB8GEOjm4UUVY3XR56HE6CyKCaO2eu5r9gqq6CY5Rbt5JpJdnEhvNErCpFRUzG3UY+ShQRGq5xkJYWFJ+9olFJND0xlUzY1Y9zBxNI0qvS1AemnJZ7Oyzx67yuR4sQ8K3KQTgZEX6BCKV/38V3LhJdfIfRsfVqENq8oLjnyOvOT0E2SXndcvyhATPz5/7n/KZVde34lzT6/D3/03zDqtewAAIABJREFUf/EurSR3s6Dy/H/+/rMyMTHW03bd7Du5zqCCGJw+84VvbfO2O+4wJUc87xDZf989Zf2O62TnnXbU6nnG9+X8i6+WCy66Wt799tfK3nvtvmDbQceDna0EO3j1cdO2x+67yAfeffai/nhca5/90rfl8qt+0fFU/MWH3iaHHXpQx/XarYDr9Mabb5Pzfna53HnPRr05wh/x3XbdIKec+AI58bgjpFBYfO5EuVyRy678hVxwydWCY8PNFs7RSSccLWeccmzbbR3X2bn5rixRjz+xWT7+qXP1fOBm7NWvfImcdcaJ2haaCwmQAAmsNgJIT3li07Q88MjjMjdb0gq3B3Gu1hNYXeBDxyTRqvrQ4UtH0osXwuYSaEshpJbrI4XGQM1sdJTX1Y9uF9MsyWSvawVdXTWR6RyKqacqvjNqb9EoRRunmM6ajqGae67rGPE/VsjLLhs2yA5Tk0xoWcELc6QFOsTGJz/zdblv40NtEUMkvuftr5PnHf7stq//9ra75f//xy9oXNGOO6wVf5H2tfgiy1WN161dI4V8bslTusMOa7UaP9nlDOh+Raab9Nnt9fXLG38rF192rZx84tGLMmm3LzeptFavC2wtyG4dHy/Kr39zh3z045+Xw5/zzLZCFsLyu/99gXzvhxfqRFDYh4456vDGH4RhFOg4vn/93Df0m4LFjssxmp2dl4/+4+cF38J0WvoV6LAOffIz/y633XHvom+B84MbiWfsscs26+C6/ad/+arcfe8DbbfHTcgH33P2NvMuehHoreL8Da95ubz41OOYm9vpouDrJEACq4LAluk52fjQo7Jl84z51jhCN1Iku5hJoohb1ImjmvhSlxQsL6icazZ6pJZw337JrBYXW013HUQhyheIdBv9qF50zWyEOM/qT4jylJ+XVNZ0DNUOovCep9OydnJSdt1pvayZGF8V3If9IEZWoEN4fOwTX9aKHSq4b37DH8jhhz5T7zrvvmejfOlr31XhDkH9Pz/0p7LnHrsuOJcQYp/4zNfkpptvl1f83sla8XMfhtaTjrSRf/7Xr8qNv75N+hVaS11I/Qp0J3BX+iJdbFIp2C0l0DEu/FG59Irr5Mv//j0JgkB+/6Wnyh+87DStrA6bQJ+bL8m/ffFb8otf3qJIOwl0U2n/nGzeslVe+YozZM9nLLzGkudlv32eIWunJns6VbjuPvVvX5cbfnWLZLMZec0fnSUvOu5I/XYCN0u/+OVv5Mtf+56gyr3XnrvJh9//Fr3RdEvyZgOfkXPe+mp51kH76suoyJ/779+VJzdt1d996L1v1psut3Qr0JPiHFX9s1/3cjn95GNZlenpTHNlEiCB1UCgXKnKI489KY8++qSUKiWJNGoRCS9mwii86aZ5kfOjG6uLr3VwJCCaHHKdw2k7FKkwR4yis7S4SjoCZCDzdWKomQQKKwuSWtTakslJOpOTQnFcNqxfr99u53NLFxdXwzkYpmMYSYGOC/Zb3/mJfP9HF6sA/8ifvVX23nOhfQLCAbaRxx5/Ul72kpPkta86a8EFf94FV8pXv/F9tTH8zw/+qYr8xZZhFeg333KnoCrezXL3vRvlzrs3NlaF+Nxlp8XtP8l94huIQw8+YJu36UagO5GOcf7r57+pFfhTXvQCOfu1L5eLL7924FSZfm9uWg/mwYcflX/97DcXfBvTSaCD/0c//jlZt3ZqRdJ/XLpQEAby3ne8QY456rBtzsGdd92v3wLNl8p6jb/8rJMb61x3w816Y1nI5/UaP2D/vRZsjxvYv/+nL+kE33Pe+iq1vPQi0HGD8o//8lV58KFH9AbiLW/4fXnR8UdRnHfzgeQ6JEACq5rA1ulZeeTRx+SJx5+Qamm+KdDtpFEzYTTQCaPp2MQtwh5u/OfNpkFOoG9TRVd7S9rYWxCnCBuL+s5zUhifkB037CI77byzTE70Vhha1SflKT64kRToqJ6jcgmB0CosHH9c1F/+9+/KpVdcL0c89xB525tf2fDawr7wiU9/TSu6iwmf5HkcVoHe7bU2Mzsvf/ePX1hgc+gkPrvZd7cC3e3r5lvukE999hvyey8+Uf3J8EQPGvs4qECHP/u7P7xQfvzTy9Xbhxu1QqGgorMTox//9DL52n/8tzzzwH3kwx94q2By7HIuuH5/duFVS+4/6YN//uHPlve98w2Sy2XVp/4vn/2GQKSf8MIj9HPSmk6Ez8g3v/1j+e+fXCIHP2t/raK7PgCdKuit4hypPbiBWOxbqOXkwn2RAAmQwPZEoFQqyZYnn5RNTzwmWx57RGrzMyY7PUZ3URe96GmPoXYC3XnOGz50m4+I6aXIPM+NT8q6nXeXdTvtIut22CDFMTYYGobrYyQF+s+vv0n++V+/pr5Z5JYnv9bvdFKS1pjTTzlWzn7tKzpOZNueBTpE2Pd/dJF89wcXym677qQV4j1221keeewJ9YSj4tnNkmTQaf2J8TH5q4+co7aL1mXT5q2yZnJiaCwuX/zKd+SCS65R8XrWi0+QP3jpqfKN//sj/d1SAj0pjM849YXyljf8QScsPb1erdbk3K9/Ty1YGMfb3vxHi4pfdwzJ8Tr7Da73d/7pq+XE445s+/74ZuPv//lLek7++s/P0W+UsCwl0FvF+fve8QZ5/nOfTXHe0xnmyiRAAqNKoFouy/zstJRmt0p5ZquUZ7dIWK3YeMZA/epoeKR2F01myYinHUB98bM5yU+uleLElBQnp/RnvkhBPozX0kgKdAgoVP1ecORh8q63v6br7E6kVyDt5StqbdlZPbuoNKK75lJLvR7ID358sdx3/0Pyylecvk0aSeu2vXbqHLQKvNTYkUv+/33sc7Lfvs+QA/ffW/7vf50nf/yHZ8odd90nv3vkCRXSG9av63htO8F4x533NdadnS/JzMyc2htwk+SCFIvFgpzzJ38sW7fO6MRFiL92yzB40HEtwQ9/xqnHNVJ32gne1vEnJ4jCo4+WyS4lBevi5vG0k47tmLDSEXyHFZKV8mQF/dbb75GP/sPnNbrrf334HDlgv/bxoUkhnpxfsZhAT4rzTpOwBz02bk8CJEACJEAC2yuBkRPoEIpItYBNBZXLN7/+9+XOu+9XwY5UEQgWxCwee/Th6seFRz25oKL8y5tulQ07rlPxuBITLXvt1LlSAt1NhL3r7o3qQb77vgcalhJMmoV3GakqqM72GofnLETnX3R120rzRZf+XL741e+ozx2JNq0RizgnwyDQ233wuxHosMD8zd99VqZnZpf827FUwspy/NFZzIPuvmVa6tsMvD/853/zd/+mdrHkddtOoD/w4CPqaccNLfb7vne+Xg49+MDlOAzugwRIgARIgARWFYGRE+hJq8XrX/1SgYcYMX6ojrcu3YgIJxIPOnAf2adloqnbH1rqwgqAxIvnP/dgFfdLLYtNqlxsm5UQ6Ig4/Pq3fqgi+KwXnyive9VZmknuPN+o7n71m9/XbxRg88HNTi/+4eQ8AFgrYHMoFHKNfeBG6qvf/IFAqKPS2hqxuL0LdOe/x3Eg7/4Nr3mZPO+wZ+ux4luFS6+8Xv7zez/TG8Z2CSvL8VeoNaUFzaFQuU+y7ZTJj4lMf/3RT3cU6Ac/c/+GOMf+V+qYloML90ECJEACJEACTzeBkRPoSUGB6vjW6Rk56IB91AO8+247axTR7XfeK1/66n9ppW+xlBd34rqp4m6PHnREBeKbBvjOkXIDDq3H6pJuSqWy/K8Pv132aWkktNTF7RJCcGOkEYJrJ/U9Xv/HL9VJiliSOej4d2vzmm7Yd/qArcTNTTcVdERHfve/L5IwDOR/vP8t26QIYdxuMjJEeqcoz07H2fp68gYIjYHe/pZXLZhP4Nguh0A/7pjnyb33PbiNFQy+9n6+fen1WLk+CZAACZAACWxvBEZaoONkwXf77re/rpE+4U5gMmbxpBOO0gpua4pFstK4lC1lexPo8B8jPxu53h9875s0H36xY73mupvkU/8GIb9z24Y17T4QrRNGd9qwo8RxJGB+xPMOFiR6jI+ZTG1YYVC5/+o3fiBrJsf1RgCTVLtl3+kD+XQJ9E7jcjcon/nCf2in291320n+6iPv6DkLvd37oHL+lW98T7uDLta1czkFuhsDEm7+7D1v0s67sDbhvfv59qUbdlyHBEiABEiABLZnAiMt0BFrh6/1991nj7bnEBM7ESO304Yd5K///B0a1N+6dFPF3Z4Eusu2xjcLreKp3bEmrTDd2hZcPne+kFM7Bywur3vV72nOOQQzWtq/622vaUwOhUi/5robZYd1a+WgA/ZunIJu2Hf6cA6zQMfYIWY/84VvycTEmPzl/3h7xwnGnY4XnvdPf/4/NN1lMXGevPlZjgo69pf00ieTkDp9Q9XpePg6CZAACZAACaxGAiMn0JNiGZFw//vP36GTQtstzidcyOcWTbLoRiRuLwI9mbDRzn6w2LEmvcywq/zZe85eNLrSsfjtbXfLqScdIz/52eWNSaIQ6x//1LmyadNW+cC736jZ2kst3bDv9KEddoGe9Kq7lBRnoWl3bEsJ6uT5xbdBb3zNy+T0U16oQr11cfGJvUwSfesb/1AQPYolOUkUn7MPvufshr8dryftO2gS5mxUnc4XXycBEiABEiCBUSAwcgIdLc4//bn/EKRUIEECFg4I8HYLYhH/9h8+J4jES0bIJdftRiRuDwI9Kd5aK9jueJc61mRVtJ0gc/twlhg0f0IDnH/45JcXpLigI+f8fHlBpbz13KCiPjdXkgsvvUa+9Z3z5OQTjxakyjz62JPy+BOb9IYL6TzOy77UB/npEujuGCCUXXOfbm8S+xHoaPL0yc98XWbn5vX93v221womIy82sXelYhbdMSIH/hvf/rGgWROWM087XjBpu9c0oFH4I81jJAESIAESGD0CIyfQcYr/6wcXyLe/+1Ot6MG60hql6C4DV710VcQN63eQn5x/uYpDt9y78SG5/Y57ZTlTXLBvZIOjugnfbqdlUJGZjL9byqbS6WYk2fodzN59zmvlsEMOaohApNggmvGxx57UCjnE6Uc//vkFAh2TF++69wFN13no4cc0Cx22DBxjtVaXzVu2tk3cSTLqZMtIrjsou3bnptMkUXxT8L//7t+02+hiXTrdfl23UeTEo6mWS1npdE0kX09Wq2E1WSy2MrnNE09u1vjExx7fpOcKPQPaLf00KnL7wecI35jgZgBV/G668vZy3FyXBEiABEiABLZXAiMp0F11MAgD+cC7z5ajnn/oNucPFc7vfP98fcCj/hcffJuEUdSIlFvpE97JWrBcIhM3F5/67Nc1ArJT5nYngY4x3X3vA/LxT56r+dgQ4K/+wzPlxacdJ6lUqhHbeMzRh8s7/+TVKsxaBXryW4vFGLvGRpjECrELKwwm8ubzOdl15/UyMTGuk0zbWTda9/l0CPRkF9GlutnCOgQBC88+rtH3nPM6vXHrZUmK82ceuI+855zXd3XTl2xgtNgkaXxGMEcDPQRwDj703jc3vg1YqpNocvyuERYq+/Sj93JmuS4JkAAJkMBqJjCSAj0pfBarGCerwS7iDh1Bb7/zPqnWaj1dE712EsXO8VX//vvu2UgzWeoN+xGZbuLlF77yHQEPJ86npiY1wWXTpi3bvKXr/AkLyYRNWUmudOABe6u15JFHn5CPfeJLKvqxvPAFz9UOqh/7xJdlenpWmx4dsP9eOlGxVaBjm/MvukogYtGcyInusbGiPs/bCEbst5sbhk4nqh92nfbZqYKO7ZMxk2e/7hVy5mnHLbCb4Pz87MKrNGveT/taxYYlpZclmUTU7QTe5P4RBfm5L39bCvm8fOQDb9VviRYT1+e89VVy0glHN17uVqAnjxORm7BXvf+db1zS9tMLA65LAiRAAiRAAtsjgZEU6DhRd969Uf7+n76onlx0BH3T616heehYbrz5Njn337+rArMfYdN6IQyjBz2ZMY5unfgmARySOfG9XtBIY8F+IKIhDj/x6a9JEIby4fe/RSeNIhUHmenIM4f3uZ1A7+U9t2eBjmsCN0I3/OoW/abhrBefIC9/yckyPl7UbwXACnYqiNZ+/Nm4wUGjJ4h83Oz9wctOkz2fseuSeHPZrHr/XZU+OfkX3+i89ew/lCOfd8g2n5FnHbSvVs8xdrd0K9CxfpIF/u0aY7WLNe3l+uC6JEACJEACJLC9EhhZgY4Tlpw41+4EdrJ8dHvSh1GgY+wQf9ffcLPss/cesmH90t1NsX6vghh+cjR+KhTyiqpcqUo6lWoIwFEW6OABb/3nvvRtTTRpt8CigzSdN77m5T1XlDHp928/9jnB5N1ul3befWz/T//yVbUutVsWmxDci0DHfs14P6s3xbhBeP+73qg9CriQAAmQAAmQwCgSGGmBjhOOivHPLrxSrr72Rk0AwQKhglb2Jx53RENcDnJxDKtA7/WYehXonfY/6gLd3ST96te3yk9+epncec9GgfcbKSuHPOsAOeuME+SA/ffuykvfyjoZz9jpPLjXF5tciwm7aGp0wSVXq5DGjd3OO+2olpYzTjm27WekV4GOMbiEH+wfk6PRo6CfSbHdHi/XIwESIAESIIFhJTDyAv2pODEU6O0pr1aB/lRcU3wPEiABEiABEiCB1UuAAv0pOLcU6BToT8FlxrcgARIgARIgARJYJQQo0J+CE0mBToH+FFxmfAsSIAESIAESIIFVQoAC/Sk4kZgsee7Xvyd33Hmf/Omb/0iQerGcy0pEBbYb38233CloTIO4v0MPPmDgQ0AO+hfO/U9x8YzddP5MviliAP/7x5fIy846SV50/FF9jeepYtfX4LgRCZAACZAACZDASBKgQF8Fp50is/+TSHb9s+OWJEACJEACJEACK0OAAn1luD6le6XI7B832fXPjluSAAmQAAmQAAmsDAEK9JXhyr2SAAmQAAmQAAmQAAmQQF8EKND7wsaNSIAESIAESIAESIAESGBlCFCgrwxX7pUESIAESIAESIAESIAE+iJAgd4XNm5EAiRAAiRAAiRAAiRAAitDgAJ9ZbhyryRAAiRAAiRAAiRAAiTQFwEK9L6wcSMSIAESIAESIAESIAESWBkCFOgrw5V7JQESIAESIAESIAESIIG+CFCg94WNG5EACZAACZAACZAACZDAyhCgQF8ZrtwrCZAACZAACZAACZAACfRFgAK9L2zciARIgARIgARIgARIgARWhgAF+spw5V5JgARIgARIgARIgARIoC8CFOh9YeNGJEACJEACJEACJEACJLAyBCjQV4Yr90oCJEACJEACJEACJEACfRGgQO8LGzciARIgARIgARIgARIggZUhQIG+Mly5VxIgARIgARIgARIgARLoiwAFel/YuBEJkAAJkAAJkAAJkAAJrAwBCvSV4cq9kgAJkAAJkAAJkAAJkEBfBCjQ+8LGjUiABEiABEiABEiABEhgZQhQoK8MV+6VBEiABEiABEiABEiABPoiQIHeFzZuRAIkQAIkQAIkQAIkQAIrQ4ACfWW4cq8kQAIkQAIkQAIkQAIk0BcBCvS+sHEjEiABEiABEiABEiABElgZAhToK8OVeyUBEiABEiABEiABEiCBvghQoPeFjRuRAAmQAAmQAAmQAAmQwMoQoEBfGa7cKwmQAAmQAAmQAAmQAAn0RYACvS9s3IgESIAESIAESIAESIAEVoYABfrKcOVeSYAESIAESIAESIAESKAvAhTofWHjRiRAAiRAAiRAAiRAAiSwMgQo0FeGK/dKAiRAAiRAAiRAAiRAAn0RoEDvCxs3IgESIAESIAESIAESIIGVIUCBvjJcuVcSIAESIAESIAESIAES6IsABXpf2LgRCZAACZAACZAACZAACawMAQr0leHKvZIACZAACZAACZAACZBAXwQo0PvCxo1IgARIgARIgARIgARIYGUIUKCvDFfulQRIgARIgARIgARIgAT6IkCB3hc2bkQCJEACJEACJEACJEACK0OAAn1luHKvJEACJEACJEACJEACJNAXAQr0vrBxIxIgARIgARIgARIgARJYGQIU6CvDlXslARIgARIgARIgARIggb4IUKD3hY0bkQAJkAAJkAAJkAAJkMDKEKBAXxmu3CsJkAAJkAAJkAAJkAAJ9EWAAr0vbNyIBEiABEiABEiABEiABFaGAAX6ynDlXkmABEiABEiABEiABEigLwIU6H1h40YkQAIkQAIkQAIkQAIksDIEKNBXhiv3SgIkQAIkQAIkQAIkQAJ9EaBA7wsbNyIBEiABEiABEiABEiCBlSFAgb4yXLlXEiABEiABEiABEiABEuiLAAV6X9i4EQmQAAmQAAmQAAmQAAmsDAEK9JXhyr2SAAmQAAmQAAmQAAmQQF8ERlqgP/DgI3Ljr2+XTZu3ShiGCnB8vCiHPPsAedZB+0gqldoGar1el5tuvl3uuGujlMuVrraJokjuuucB+c0td8r0zJzg39j3mslxed7hz5K999q948kLglB+duFVsnV6Rl565otkcnK84zatK2ydnpXrf/EbefiRxyUIAkmn07LDuik56ohDZOeddmy7v36Ot9PA+uXRz/grlaqe43vue7Bxvnzfl9122SBHHnGITK2Z2Ga4/Y6v03G71+fnS3Le+Vcq/zNPP07y+dySm/76N3fI9Tf8Ro479nly0AF7d/s2XI8ESIAESIAESGA7JTCyAv26X9wsN99yp542CKVMxlfhXKvV9Xe77LxeTj7xaCkUmuIJYvWiS6+Vhx5+TNfJ5bISx3Fjm/32eYaKKN9PNy4HCOtLLr9ONj7wO/0dxCFex+8hkrG02y55PWHdK6/+pdx97wM6nn4E+hNPbpELLr5GSqWy3hxksxmp1wO9McF4jjn6cDlw/70WXMblclXOv+gqwbbuePGzWq0tyqjT56BfHv2Mf/OWaTn/oqtlbq7UGL/neXq+cK7bHXe/4+t03O51ML34smvlkUef0JujTgL9jrvul2uuvVGvFwr0bilzPRIgARIgARLYvgmMpEC/b+PDctkV10sUxXL4c54phx16YKNaftfdG+Xqa28SiPGDn7WfvOCowxpn2In6YrEgp7zoaNlpww76mtsGYvfYow+Xgw5sVjld9TOTycixRx8m+++3p24DgXjTzXfIjb++TUU+3ufZz9x3m6tpvlSWSy67Th597El9rR+BjiryTy+4Up7ctFV223WD3njg5gLH+PPrfy133Hm/jI0V5CWnHy9rEhXlK67+pdxx530yNlaU004+RnbcYUrHAHF5yeXXq9h/7mHP0m8Bul364dHP+N03DhjrurVr5OQXHd2olmN/uGl6+HePy1ixoCJ5ampSD6Gf8XV77Js3T8tFl10r09OzuslSAh3Xxw2/+q385rd36bWChQK9W9JcjwRIgARIgAS2bwIjKdAvvvRauff+h2TvPXeTU056wTZn0Ik0iNWzzjhBisW8zM7Oy49+epmUShU5/oXPlwOs0HYbu23W77hWXnLG8QJBjgr1eedfIY8/sVkOO/QgOeJ5B2/zXk4Ew2ICoYhqPhYIzN/edrcKeOwHlV8I+X4E+p13b5QrrrpBj+P3XnyiTEyMNcaRFLLJMcJO8uOfXi4Qs0sd7447mONFRb7T0i+Pfsb/2OOb1BIUxyJnnHrsNhYeZzPBcR75/EPkOYccOND5WurY8Y3DL2+8VW6/8z79xsKdy8UEOm4qrrrmV4KxYXHrU6B3usL4OgmQAAmQAAmsDgIjJ9Bhb7jwkp8LBNwxRx/W1tMLC8uFl1yjthdnJ4G95LIrfqHebyfak5cAqqI/PO8yrXaeceoLtbo+OzcvP7vwavU+n3TCUbL7bjttc9VAtMG+0irW3O+xAfZ14P57yzXX3bhgTN1egu6G5ID995ITXvj8bTb77W33qI1iw/p1cubpuLnwtWKPcWHB71BhTy5ufBD7v/fiE7TK3mnpl0c/48f5uubam2R8rLiojQTfKuBcu29K+h1fp+P++XU3yS233q2r7bPX7rJ+/TrBtzHtBPrMzJz88LxLBVYYfFPz/MOfrTdqmCdBgd6JNF8nARIgARIggdVBYOQEejen7f6ND6uFI5ttCnRnb4FFBIK1dUGlGRP/ehVSrvLeKtZgm4HIhH1kw/odVEi23jR0cyy4IfnJz66QJzdtUZ95OxtNP/t2lX/wOP2UYxuV/27GtNQ6rTzgl1+J8aOSjfOFG5FWK1Mv4+s0wRP7+uWNv9XJwRDbuMFb7KYM6+KbmsuvukGtUPvv+wz1y/dzXQ16Hrg9CZAACZAACZDA00eAAr0N+4su+bnAp54Un66Kiwoo/MztltaKbKfTCg84xCcmQCKdAxXSxZZ+RDT2NT9flh/+5FKBl/3E44/QCamti6vawoJy6knHtK30u20gIK+74Tdy3/0PqY3n5BOPkj1237nToXb1ejseyz1+NxAIc3y7EceRnHTiUbLnHrt2HGMv52upnS0l0Jfrxq/jwXAFEiABEiABEiCBoSVAgd5yamAnuPb6myWV8uTE449UnzqWbsR3N+sk3+7n1/1abrn1Lp2wiaq8m4TZ7mrpV6B3I767Wce9PzzrWOBnP/G4I/UmZrmWdjy6GVs36yTHCPsI/On4VgFpPbAkJZN3FjueXs4XBfpyXRXcDwmQAAmQAAmMHgEK9MQ5T0baHfys/eUFRz2n8Wo34rubddwOf/HLWxoxj0cfeag8+5n7LXn1Pd0C/Z57H5RrrrtJx+hiCuFVR/pMazxjPx+jxXh0I767WceNKRlzCN88JpAi5aXT0uv5okDvRJSvkwAJkAAJkAAJLEaAAt2SgfcZSRvwJsNuAuGZrKp2I767WQcVaPjZb739Ho12h"
# def analysis_base64(src):
#     result = re.search("data:image/(?P<ext>.*?);base64,(?P<data>.*)", src, re.DOTALL)
#     if result:
#         ext = result.groupdict().get("ext")
#         data = result.groupdict().get("data")
#     else:
#         raise Exception("Do not parse!")
#     img = base64.urlsafe_b64decode(data)
#     return img
# analysis_base64(src)

# df = pd.read_csv(r'C:\Users\wb\Desktop\安克创新2.0.csv')
# Note=open(r'C:\Users\wb\Desktop\text.txt',mode='a',encoding='utf-8')
# for i in df.index:
#     div = "<div>-------------------------------------------------------%s--------------------------------------------</div></br>"%(i+1)
#     Note.write(div)
#     text = df.loc[i,"字段"]
#     img = df.loc[i,"字段1"]
#     Note.write(text)
#     Note.write("</br></br>")
#     Note.write(img)
# Note.close()

os.rename(r'C:\Users\wb\Desktop\text.txt',r'C:\Users\wb\Desktop\text.html')
