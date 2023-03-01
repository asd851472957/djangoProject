import os
import sys
import pandas as pd

input('Press Enter to exit...')
def app_path():
    if hasattr(sys, 'frozen'):
        return os.path.dirname(sys.executable)  # 使用pyinstaller打包后的exe目录
    return os.path.dirname(__file__)  # 没打包前的py目录


import winreg
def desktop_path():
    key = winreg.OpenKey(winreg.HKEY_CURRENT_USER,
                         r'Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders')
    desktop = winreg.QueryValueEx(key, "Desktop")[0]
    print(desktop)
    return desktop

from docx import Document
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None
desktop_path = desktop_path()
doc = Document(desktop_path + "\日本海外仓箱标\日本海外仓箱标模板.docx")
df = pd.read_excel(desktop_path + "\日本海外仓箱标\需生成的数据文件.xlsx")
df_data = pd.DataFrame()
for i in df.index:
    xiangshu = int(df.loc[i, "箱数"])
    sku = df.loc[i, "品名"]
    num = int(df.loc[i, "单箱数量"])
    zimu = df.loc[i, "字母标签"]
    code = df.loc[i, "箱子编号"]
    if xiangshu > 1:
        df_tmp = pd.DataFrame({"品名": sku, "箱子编号": code, '字母标签': zimu, "单箱数量": num}, index=[0])
        for n in range(0, xiangshu):
            df_data = df_data.append(df_tmp)
    else:
        df_tmp = pd.DataFrame({"品名": sku, "箱子编号": code, '字母标签': zimu, "单箱数量": num}, index=[0])
        df_data = df_data.append(df_tmp)

df_data = df_data.reset_index()
sku_list = df_data['品名'].tolist()
num_list = df_data['单箱数量'].tolist()
zimu_list = df_data['字母标签'].tolist()
code_list = df_data['箱子编号'].tolist()
list1 = doc.paragraphs
for i in list1:
    delete_paragraph(i)
    if (len(doc.paragraphs)) == len(sku_list):
        break

doc.save(desktop_path + "\日本海外仓箱标\日本海外仓箱标.docx")
doc = Document(desktop_path + "\日本海外仓箱标\日本海外仓箱标.docx")
children = doc.element.body.iter()
count = 0  # 写一个count是为了，可以定位是哪个文本框，因为我用索引失败了
for child in children:
    # 通过类型判断目录
    if child.tag.endswith('txbx'):
        count += 1
        if count == 3:
            for ci in child.iter():
                if ci.tag.endswith('main}r'):
                    if ci.text == '型号：':
                        ci.text = ''
                    if ci.text == '字母区分：':
                        ci.text = ''
                    if ci.text == '数量：':
                        ci.text = ''
                    if ci.text == '箱子编号：':
                        ci.text = ''
doc.save(desktop_path + "\日本海外仓箱标\日本海外仓箱标.docx")
doc = Document(desktop_path + "\日本海外仓箱标\日本海外仓箱标.docx")
sku_dict = {}
for i in range(1, (len(sku_list) * 2) + 1):
    if i == 1:
        sku_dict[1] = 1
    elif i % 2 == 0:
        sku_dict[i] = sku_dict[i - 1]
    else:
        sku_dict[i] = sku_dict[i - 1] + 1
m = 1
n = 0
children = doc.element.body.iter()
for child in children:
    # 通过类型判断目录
    if child.tag.endswith('txbx'):
        for ci in child.iter():
            if ci.tag.endswith('main}r'):
                if ci.text == '型号：':
                    # print(ci.text)
                    ci.text = "型号：%s" % sku_list[n]
                if ci.text == '字母区分：':
                    # print(ci.text)
                    ci.text = "字母区分：%s" % zimu_list[n]
                if ci.text == '数量：':
                    # print(ci.text)
                    ci.text = "数量：%s" % num_list[n]
                if ci.text == '箱子编号：':
                    # print(ci.text)
                    ci.text = "箱子编号：\n%s" % code_list[n]
                    n += 1
            print("正在生成")

doc.save(desktop_path + "\日本海外仓箱标\日本海外仓箱标生成完毕.docx")


input('Press Enter to exit...')